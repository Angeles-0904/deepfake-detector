"""
Script para convertir documentacion_completa_proyecto.md a Word (.docx)
con formato profesional: tablas, títulos, negritas, etc.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

MD_PATH = Path("documentacion/documentacion_completa_proyecto.md")
DOCX_PATH = Path("documentacion/Documentacion_Completa_DeepFake_Detector.docx")


def parse_markdown_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    
    # ─── Estilos base ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # ─── Configurar márgenes ─────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ─── Leer el markdown ────────────────────────────────────────────────────
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")
    
    # Variables de estado
    in_table = False
    table_rows = []
    in_code_block = False
    code_buffer = []
    in_list = False
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # ── Código bloque ────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code_block:
                # Cerrar bloque de código
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # ── Tablas Markdown ──────────────────────────────────────────────────
        if "|" in stripped and stripped.startswith("|") and stripped.endswith("|"):
            if not in_table:
                in_table = True
                table_rows = [stripped]
            else:
                table_rows.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                # Procesar tabla acumulada
                _add_table(doc, table_rows)
                in_table = False
                table_rows = []
                doc.add_paragraph()  # espacio después de tabla
        
        # ── Saltar línea separadora de tablas (| --- | --- |) ────────────────
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            i += 1
            continue
        
        # ── Títulos ──────────────────────────────────────────────────────────
        if stripped.startswith("# ") and not stripped.startswith("```"):
            p = doc.add_heading(stripped[2:], level=1)
            i += 1
            continue
        
        if stripped.startswith("## ") and not stripped.startswith("```"):
            p = doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        
        if stripped.startswith("### ") and not stripped.startswith("```"):
            p = doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        
        if stripped.startswith("#### ") and not stripped.startswith("```"):
            p = doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        
        # ── Saltos de sección (---) ──────────────────────────────────────────
        if stripped == "---":
            doc.add_paragraph("─" * 50)
            i += 1
            continue
        
        # ── Línea vacía ──────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue
        
        # ── Párrafo normal con formato inline ────────────────────────────────
        p = doc.add_paragraph()
        _add_inline_text(p, stripped)
        i += 1
    
    # Si queda tabla sin cerrar
    if in_table and table_rows:
        _add_table(doc, table_rows)
    
    # ─── Guardar ─────────────────────────────────────────────────────────────
    doc.save(docx_path)
    print(f"[OK] Documento Word creado: {docx_path}")
    print(f"       Tamanio: {docx_path.stat().st_size / 1024:.1f} KB")


def _add_inline_text(paragraph, text: str):
    """
    Procesa formato inline: **negrita**, *cursiva*, `código`, y enlaces [texto](url).
    """
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\))'
    last_end = 0
    for match in re.finditer(pattern, text):
        start, end = match.start(), match.end()
        
        # Texto antes del match
        if start > last_end:
            paragraph.add_run(text[last_end:start])
        
        full, bold_text, italic_text, code_text, link_text, link_url = match.groups()
        
        if bold_text:
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif italic_text:
            run = paragraph.add_run(italic_text)
            run.italic = True
        elif code_text:
            run = paragraph.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif link_text and link_url:
            run = paragraph.add_run(f"{link_text} ({link_url})")
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
            run.underline = True
        
        last_end = end
    
    # Texto restante
    if last_end < len(text):
        remaining = text[last_end:]
        # También procesar listas de items con - al inicio
        remaining = re.sub(r'^- ', '• ', remaining)
        paragraph.add_run(remaining)


def _add_table(doc, rows: list):
    """
    Convierte filas markdown a tabla de Word.
    La primera fila se toma como encabezado.
    """
    # Limpiar y parsear filas
    parsed_rows = []
    for row_text in rows:
        # Quitar | inicial y final, dividir por |
        cells = [c.strip() for c in row_text.strip("|").split("|")]
        # Si la fila es separador (|---|), omitir
        if all(re.match(r'^[\s\-:]+$', c) for c in cells):
            continue
        parsed_rows.append(cells)
    
    if not parsed_rows:
        return
    
    # Crear tabla
    num_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_cells in enumerate(parsed_rows):
        for col_idx in range(num_cols):
            cell_text = row_cells[col_idx] if col_idx < len(row_cells) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            
            # Primera fila = encabezado en negrita
            if row_idx == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    parse_markdown_to_docx(MD_PATH, DOCX_PATH)
