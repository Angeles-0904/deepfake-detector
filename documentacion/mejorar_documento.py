"""
Script para crear una versión mejorada del documento samely.docx
- Antecedentes Internacionales mejorados (del Marco Teórico)
- Metodología actualizada (según implementación real)
- Secciones incorrectas corregidas
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import os

# ============================================================
# CONTENIDO MEJORADO
# ============================================================

ANTECEDENTES_INTERNACIONALES = [
    # Los antecedentes del Marco Teórico son más completos y recientes
    (
        "Abdul-Hafiz & Sari (2025) desarrollaron un framework robusto y explicable para detectar deepfakes en imágenes utilizando transfer learning con mecanismos de atención (CBAM). Emplearon 5 redes neuronales convolucionales en paralelo, las combinaron con XGBoost y aplicaron Grad-CAM para generar mapas de calor. Con ello resolvieron el problema de las \"cajas negras\" en detectores de deepfakes, añadiendo explicabilidad visual. Alcanzaron una precisión del 97.25% en un dataset de 140,000 imágenes, notablemente similar al dataset utilizado en la presente investigación."
    ),
    (
        "Rakesh Kumar (2025) propuso un sistema de detección de deepfakes basado en redes neuronales convolucionales con mecanismos de atención, integrando técnicas de Inteligencia Artificial Explicable (XAI) como Grad-CAM, LIME y SHAP para proporcionar transparencia en las decisiones del modelo. El autor comparó múltiples arquitecturas pre-entrenadas (EfficientNet, ResNet y VGG) y evaluó su sistema en un dataset de 1,000 imágenes, alcanzando una precisión del 94.7% y un AUC de 0.967. Este trabajo demuestra que superar el 90% de acierto es perfectamente alcanzable con las técnicas adecuadas de transfer learning y explicabilidad."
    ),
    (
        "Mohit Kumar (2025) desarrolló un detector de deepfakes centrado en la robustez del modelo, evaluando su rendimiento bajo condiciones adversas como compresión JPEG, inyección de ruido y ataques adversariales. Utilizando arquitecturas pre-entrenadas (EfficientNet, ResNet y VGG) en combinación con técnicas de explicabilidad (Grad-CAM, LIME y SHAP), el autor demostró que es posible mantener una alta precisión (94.7%) incluso cuando las imágenes son sometidas a degradaciones propias de la transmisión en redes sociales. Este trabajo inspira directamente la evaluación del presente sistema con imágenes comprimidas."
    ),
    (
        "Aribe Jr. (2025) propuso un enfoque híbrido que fusiona características forenses tradicionales con aprendizaje profundo, con el objetivo de mejorar tanto la robustez como la explicabilidad de la detección de deepfakes. Su metodología integra el análisis de ruido PRNU, artefactos JPEG y coeficientes DCT con una arquitectura que combina ResNet-50 y Vision Transformer. Utilizó Grad-CAM para generar mapas de calor, complementándolos con mapas de calor forenses para una explicabilidad dual. El resultado más relevante es un overlap del 82% entre los mapas de calor generados por Grad-CAM y las regiones realmente manipuladas."
    ),
    (
        "Lipianina-Honcharenko et al. (2025) implementaron un método de conjunto (ensemble) basado en cinco redes neuronales convolucionales (ResNet50, EfficientNetB0, Xception, InceptionV3 y FaceNet) con selección de \"golden frames\". Las predicciones de las cinco CNNs fueron combinadas mediante un meta-modelo XGBoost, logrando una precisión del 91.14%. La selección de golden frames mejoró la precisión de Xception en un 13.9% y la de ResNet50 en un 11.1%, demostrando la importancia de optimizar los recursos computacionales."
    ),
    (
        "Raikwar et al. (2025) desarrollaron un modelo híbrido que combina redes neuronales convolucionales con transformers para la detección de deepfakes, resolviendo la necesidad de mejorar la precisión mediante la combinación de capacidades locales (CNN) y globales (transformers). Obtuvieron una precisión del 94.8% y un F1-score del 93.9%."
    ),
    (
        "Singh & Kumar (2025) propusieron una arquitectura CNN-RNN para la detección de deepfakes en imágenes, combinando la extracción de características espaciales con el modelado de dependencias temporales, capturando tanto patrones espaciales (texturas, bordes) como dependencias secuenciales."
    ),
    (
        "Wang et al. (2020) desarrollaron un detector \"universal\" de imágenes generadas por redes neuronales convolucionales, evaluando su rendimiento en 11 generadores diferentes (ProGAN, StyleGAN, BigGAN, entre otros). Demostraron que es posible detectar imágenes sintéticas independientemente de la arquitectura generadora utilizada, abordando el problema de la falta de generalización."
    ),
    (
        "Selvaraju et al. (2017) introdujeron Grad-CAM (Gradient-weighted Class Activation Mapping), una técnica de explicabilidad visual que genera mapas de calor a partir de los gradientes de la última capa convolucional de una CNN. Grad-CAM se ha convertido en el estándar de facto para explicabilidad en redes neuronales convolucionales, siendo class-discriminative y permitiendo identificar las regiones que influyen en la decisión del modelo."
    ),
    (
        "Rossler et al. (2019) introdujeron FaceForensics++, un dataset de referencia para detección de deepfakes que contiene 1,000 videos reales y 4,000 videos manipulados con cuatro técnicas diferentes (Deepfakes, Face2Face, FaceSwap y NeuralTextures), ofreciendo tres niveles de compresión (c0, c23, c40). FaceForensics++ se ha convertido en el benchmark estándar en la literatura de detección de deepfakes."
    ),
    (
        "Li et al. (2020) introdujeron Celeb-DF v2, un dataset de deepfakes de alta calidad con 590 videos reales de celebridades y 5,639 videos deepfake generados con una técnica avanzada de intercambio facial, ofreciendo deepfakes mucho más difíciles de detectar que datasets anteriores."
    ),
    (
        "Tolosana et al. (2020) publicaron una revisión exhaustiva (survey) del estado del arte en detección de deepfakes y manipulación facial, cubriendo técnicas de generación, detección, datasets y desafíos abiertos. Su survey es uno de los más citados en la literatura de deepfakes."
    ),
    (
        "Guarnera, Giudice y Battiato (2020) desarrollaron la investigación DeepFake Detection by Analyzing Convolutional Traces. Su objetivo fue identificar huellas forenses ocultas producidas por los procesos convolucionales utilizados para generar rostros sintéticos. La metodología correspondió a un diseño experimental basado en procesamiento digital de imágenes y aprendizaje automático. Los autores aplicaron análisis de frecuencias y extracción de características estadísticas sobre los canales de color."
    ),
    (
        "De Lima et al. (2020) realizaron el estudio Deepfake Detection Using Spatiotemporal Convolutional Networks. Su objetivo fue mejorar la detección de videos deepfake mediante el análisis conjunto de información espacial y temporal, superando las limitaciones de los métodos que clasificaban cada fotograma de forma independiente. La investigación siguió una metodología experimental y comparativa."
    ),
    (
        "Tariq et al. (2025) desarrollaron From Prediction to Explanation: Multimodal, Explainable, and Interactive Deepfake Detection Framework for Non-Expert Users. Su objetivo fue transformar la clasificación de deepfakes en un proceso explicable y comprensible para usuarios no especializados, como periodistas, peritos y verificadores de información. La metodología fue experimental y de desarrollo tecnológico."
    ),
]

METODOLOGIA_MEJORADA = {
    "tipo_enfoque": """La investigación fue de enfoque cuantitativo, debido a que empleó datos numéricos, métricas de desempeño y procedimientos estadísticos para evaluar la capacidad del modelo de inteligencia artificial en la detección de imágenes reales y deepfakes. Hernández-Sampieri y Mendoza (2018) señalan que el enfoque cuantitativo permite medir variables, contrastar hipótesis y analizar resultados mediante procedimientos estadísticos y numéricos.

Por su finalidad, el estudio fue de tipo aplicado, puesto que utilizó conocimientos de visión por computadora, redes neuronales convolucionales, transfer learning e inteligencia artificial explicable para desarrollar una solución tecnológica orientada a un problema concreto. Según Ñaupas et al. (2018), la investigación aplicada busca emplear conocimientos científicos para resolver necesidades específicas de la sociedad.

El nivel fue explicativo-predictivo y tecnológico. Fue explicativo porque analizó las características visuales que influyen en la diferenciación entre imágenes auténticas y manipuladas; predictivo porque el modelo estimó la probabilidad de que una imagen corresponda a la clase real o deepfake; y tecnológico porque culminó en un prototipo funcional integrado en una interfaz web desarrollada con Streamlit.""",
    
    "disenio": """Se utilizó un diseño experimental computacional. Fue experimental porque se modificaron deliberadamente condiciones del entrenamiento, como la arquitectura del modelo, la tasa de aprendizaje, el número de épocas, el congelamiento de capas y el aumento de datos, para observar su efecto en el rendimiento del modelo.

El procedimiento se organizó en seis fases: (1) selección y preparación del conjunto de datos, (2) preprocesamiento de imágenes (redimensionamiento, normalización y aumento de datos), (3) entrenamiento del modelo DenseNet-121 con transfer learning, (4) evaluación del desempeño mediante métricas de clasificación, (5) generación de explicaciones con Grad-CAM, e (6) implementación del prototipo web. No se manipularon personas ni se aplicaron tratamientos a seres humanos; la unidad de análisis estuvo constituida por cada imagen digital.""",
    
    "variables": """La variable independiente fue la configuración del sistema de detección basado en la red neuronal convolucional DenseNet-121 con transfer learning y Grad-CAM. La variable dependiente fue el desempeño del sistema en la detección explicable de deepfakes, medido a través de las métricas de accuracy, precisión, recall, F1-score, AUC-ROC y tiempo de inferencia.""",
    
    "poblacion_muestra": """La población estuvo conformada por 140,000 imágenes faciales del dataset \"140k Real and Fake Faces\" (disponible en Kaggle), organizadas en dos clases balanceadas: 70,000 imágenes reales y 70,000 imágenes deepfake o sintéticas.

Se utilizó una muestra censal, empleando la totalidad de imágenes válidas disponibles (140,000). Las imágenes se dividieron mediante muestreo aleatorio estratificado: 100,000 para entrenamiento (71.4%), 20,000 para validación (14.3%) y 20,000 para prueba (14.3%). La estratificación garantizó una proporción equilibrada entre clases (50% reales, 50% fake). Se verificó que imágenes de una misma fuente no aparecieran en más de un conjunto, para prevenir fuga de información.""",
    
    "tecnicas_instrumentos": """La técnica principal fue la observación computacional, complementada con análisis experimental y pruebas funcionales. Los instrumentos incluyeron: ficha de registro del dataset, reporte de métricas de clasificación, matriz de confusión, curvas ROC, mapas Grad-CAM y registro de tiempos de inferencia.

Para el entrenamiento se utilizó PyTorch 2.x como framework de deep learning, junto con Torchvision para la carga del modelo preentrenado, Scikit-learn para las métricas de evaluación, Matplotlib y Seaborn para las visualizaciones, y OpenCV para el procesamiento de imágenes.""",
    
    "procedimiento": """El procedimiento se ejecutó en la plataforma Kaggle con aceleración GPU Tesla T4 (14.6 GB de memoria VRAM), siguiendo las siguientes etapas:

Primera etapa - Preparación del dataset: Se descargó el dataset \"140k Real and Fake Faces\" desde Kaggle, que contiene 140,000 imágenes faciales (70,000 reales y 70,000 generadas por IA). Se copiaron las imágenes a la estructura de directorios de trabajo, organizándolas en carpetas separadas por clase (real/fake) y por conjunto (train/validation/test).

Segunda etapa - Preprocesamiento: Se aplicó redimensionamiento a 224×224 píxeles (tamaño de entrada de DenseNet-121). Se normalizaron los valores de píxeles utilizando la media y desviación estándar de ImageNet. Al conjunto de entrenamiento se le aplicó aumento de datos: rotaciones aleatorias (±15 grados) y volteo horizontal aleatorio.

Tercera etapa - Arquitectura del modelo: Se implementó DenseNet-121 con pesos preentrenados en ImageNet. Se congelaron los parámetros de las capas iniciales y se reemplazó el clasificador por una nueva cabeza compuesta por: una capa fully connected de 1024 a 512 neuronas con activación ReLU, dropout de 0.3, y una capa final de 512 a 2 neuronas (clases: FAKE, REAL). En total, el modelo contó con 7,479,682 parámetros, de los cuales 1,011,778 (13.5%) fueron entrenables.

Cuarta etapa - Entrenamiento: Se utilizó el optimizador Adam con dos tasas de aprendizaje diferenciadas: 1×10⁻⁴ para el clasificador y 1×10⁻⁵ para las capas de fine-tuning. Se empleó una función de pérdida de entropía cruzada (CrossEntropyLoss) y un scheduler ReduceLROnPlateau para reducir la tasa de aprendizaje cuando la pérdida de validación dejara de mejorar. Se entrenó por 14 épocas, aplicando early stopping con paciencia de 7 épocas.

Quinta etapa - Evaluación: El modelo alcanzó un accuracy del 97.06%, un AUC de 0.9966 y un F1-score macro de 0.97 en el conjunto de prueba independiente de 20,000 imágenes. Se generaron la matriz de confusión, las curvas ROC y los mapas Grad-CAM para visualizar las regiones que influyeron en la clasificación.

Sexta etapa - Implementación del prototipo: Se desarrolló una aplicación web interactiva utilizando Streamlit, que permite al usuario cargar una imagen y obtener: (a) la clasificación como REAL o FAKE con su nivel de confianza, (b) el mapa de calor Grad-CAM superpuesto, (c) explicaciones textuales sobre las regiones analizadas y (d) pruebas de robustez frente a compresión JPEG.""",
    
    "analisis_datos": """Los resultados fueron procesados con Python, PyTorch, Scikit-learn, Pandas, Matplotlib y Seaborn. Se calcularon las siguientes métricas en el conjunto de prueba:
- Accuracy: 97.06%
- Precision (macro): 0.97
- Recall (macro): 0.97
- F1-score (macro): 0.97
- AUC-ROC: 0.9966
- Matriz de confusión: se analizaron los falsos positivos y falsos negativos.
- Tiempo de inferencia: se midió el tiempo de procesamiento por imagen.

Para cada imagen de prueba seleccionada, se generó un mapa Grad-CAM que permitió observar las regiones faciales que influyeron en la decisión del clasificador. Se consideró satisfactorio el modelo al superar ampliamente el umbral establecido de F1-score y AUC ≥ 0.85.""",
    
    "eticas": """Se utilizaron datasets públicos con fines académicos, respetando sus licencias y condiciones de uso (dataset \"140k Real and Fake Faces\" de Kaggle, bajo licencia de uso académico). No se emplearon imágenes privadas sin consentimiento. Los resultados se presentan de manera transparente, reconociendo posibles errores, sesgos y limitaciones. El sistema es descrito como herramienta de apoyo y no como sustituto de un peritaje digital especializado."""
}

# ============================================================
# CREACIÓN DEL DOCUMENTO MEJORADO
# ============================================================

def create_improved_document():
    # Cargar el documento original
    source_path = "C:/Users/SHAMELY/Documents/Downloads/samely.docx"
    dest_path = "D:/IX/detectorIA/documentacion/PROYECTO_INVESTIGACION_MEJORADO.docx"
    
    doc = Document(source_path)
    
    # Mapeo de párrafos relevantes (según el índice leído)
    # El documento tiene párrafos en este orden:
    # 35: "Antecedentes del proyecto" (título)
    # 37: "Antecedentes internacionales" (subtítulo)
    # 38-48: antecedentes internacionales (pares)
    # 50: "Antecedentes nacionales"
    # 51-59: antecedentes nacionales
    # 97: "Metodología de investigación"
    # 100-133: secciones de metodología
    
    paragraphs = doc.paragraphs
    
    # --- MEJORAR ANTECEDENTES INTERNACIONALES ---
    # Encontrar el párrafo "Antecedentes internacionales"
    for i, p in enumerate(paragraphs):
        if "Antecedentes internacionales" in p.text and i > 30:
            # Los siguientes párrafos con estilo Normal (índices pares) son los antecedentes
            # Los reemplazamos con los mejorados
            ant_idx = 0
            j = i + 1
            while j < len(paragraphs) and ant_idx < len(ANTECEDENTES_INTERNACIONALES):
                p = paragraphs[j]
                if p.style.name == 'Normal' and p.text.strip():
                    # Reemplazar el texto completo del párrafo
                    # Limpiar runs existentes
                    for run in p.runs:
                        run.text = ''
                    # Escribir nuevo texto en el primer run
                    if p.runs:
                        p.runs[0].text = ANTECEDENTES_INTERNACIONALES[ant_idx]
                    else:
                        p.add_run(ANTECEDENTES_INTERNACIONALES[ant_idx])
                    ant_idx += 1
                    j += 1
                elif "Antecedentes nacionales" in p.text or "Antecedentes internacionales" in p.text:
                    j += 1
                else:
                    j += 1
            break
    
    # --- MEJORAR METODOLOGÍA ---
    # Encontrar y reemplazar cada sección de la metodología
    
    # Mapeo de títulos de sección -> contenido mejorado
    metodologia_sections = {
        "Tipo y enfoque de investigación": METODOLOGIA_MEJORADA["tipo_enfoque"],
        "Diseño de investigación": METODOLOGIA_MEJORADA["disenio"],
        "Variables de investigación": METODOLOGIA_MEJORADA["variables"],
        "Operacionalización de variables": "",  # Mantener original, está bien
        "Población y muestra": METODOLOGIA_MEJORADA["poblacion_muestra"],
        "Técnicas e instrumentos": METODOLOGIA_MEJORADA["tecnicas_instrumentos"],
        "Procedimiento": METODOLOGIA_MEJORADA["procedimiento"],
        "Análisis de datos": METODOLOGIA_MEJORADA["analisis_datos"],
        "Consideraciones éticas": METODOLOGIA_MEJORADA["eticas"],
    }
    
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        for section_title, new_content in metodologia_sections.items():
            if section_title in text and new_content:
                # El siguiente párrafo con estilo Normal es el contenido a reemplazar
                j = i + 1
                if j < len(paragraphs):
                    next_p = paragraphs[j]
                    if next_p.style.name in ['Normal', 'List Paragraph']:
                        for run in next_p.runs:
                            run.text = ''
                        if next_p.runs:
                            next_p.runs[0].text = new_content
                        else:
                            next_p.add_run(new_content)
                break
    
    # --- CORREGIR SECCIONES INCORRECTAS (IoT/Raspberry Pi) ---
    # Buscar secciones de "Impactos" y reemplazar contenido incorrecto
    impactos_ciencia = """El proyecto generó impactos en ciencia y tecnología al aportar evidencia sobre la aplicación de redes neuronales convolucionales con transfer learning e inteligencia artificial explicable en la detección de deepfakes en imágenes. En el ámbito científico, permitió ampliar el conocimiento sobre los patrones visuales asociados con contenidos manipulados y sobre el comportamiento de la arquitectura DenseNet-121 en esta tarea específica.

En el ámbito tecnológico, se desarrolló un prototipo web funcional (Streamlit) capaz de clasificar imágenes y visualizar las regiones que influyen en la decisión del modelo mediante Grad-CAM. Esta solución puede servir como base para futuras herramientas de análisis forense digital, verificación de contenidos y prevención de la desinformación. El uso de tecnologías abiertas como Python, PyTorch, OpenCV y Streamlit favorece la reproducibilidad y democratización del acceso a estas herramientas."""

    impactos_economicos = """La propuesta conlleva una implicación económica al demostrar que es técnica y económicamente viable desarrollar un sistema de detección de deepfakes funcional utilizando hardware accesible. El modelo puede ejecutarse en CPUs convencionales con un tiempo de inferencia razonable, permitiendo su uso sin necesidad de GPUs especializadas. Esto reduce significativamente la barrera de entrada para instituciones educativas, pequeñas organizaciones y usuarios individuales que requieran verificar la autenticidad de contenido digital."""

    impactos_sociales = """La detección de deepfakes se ha convertido en una necesidad esencial en la sociedad digital actual. Los deepfakes pueden utilizarse para difundir desinformación, realizar suplantaciones de identidad, afectar la reputación de las personas y erosionar la confianza en las evidencias audiovisuales. Esta investigación contribuye a mitigar estos riesgos al proporcionar una herramienta accesible que permite identificar contenido manipulado y comprender las razones detrás de cada clasificación, fortaleciendo así la alfabetización digital y la capacidad de los ciudadanos para evaluar críticamente el contenido visual."""

    impactos_ambientales = """El impacto ambiental de esta investigación se relaciona principalmente con el consumo energético asociado al entrenamiento del modelo en GPUs. Para minimizar este impacto, se optimizó el proceso de entrenamiento mediante la plataforma Kaggle (que utiliza hardware eficiente), se redujo el número de épocas mediante early stopping y se aplicó transfer learning para disminuir el tiempo de entrenamiento. El prototipo resultante puede ejecutarse en hardware de consumo general, minimizando su huella de carbono durante la fase de inferencia."""

    # Reemplazar secciones de impactos
    for i, p in enumerate(paragraphs):
        text = p.text.strip() if p.text else ""
        
        if "Impactos en Ciencia y Tecnología" in text:
            j = i + 1
            if j < len(paragraphs):
                for run in paragraphs[j].runs:
                    run.text = ''
                if paragraphs[j].runs:
                    paragraphs[j].runs[0].text = impactos_ciencia
                else:
                    paragraphs[j].add_run(impactos_ciencia)
        
        if "Impactos económicos" in text and not "Ciencia" in text[:50]:
            j = i + 1
            if j < len(paragraphs):
                for run in paragraphs[j].runs:
                    run.text = ''
                if paragraphs[j].runs:
                    paragraphs[j].runs[0].text = impactos_economicos
                else:
                    paragraphs[j].add_run(impactos_economicos)
        
        if "Impactos sociales" in text:
            j = i + 1
            if j < len(paragraphs):
                for run in paragraphs[j].runs:
                    run.text = ''
                if paragraphs[j].runs:
                    paragraphs[j].runs[0].text = impactos_sociales
                else:
                    paragraphs[j].add_run(impactos_sociales)
        
        if "Impactos ambientales" in text:
            j = i + 1
            if j < len(paragraphs):
                for run in paragraphs[j].runs:
                    run.text = ''
                if paragraphs[j].runs:
                    paragraphs[j].runs[0].text = impactos_ambientales
                else:
                    paragraphs[j].add_run(impactos_ambientales)
    
    # CORREGIR sección Recursos (eliminar contenido IoT/Raspberry Pi)
    recursos_nuevo = """Para la ejecución del presente proyecto de investigación se requirieron los siguientes recursos, organizados en equipos de cómputo, software y datos.

Equipos:
- Una (1) laptop con procesador Intel Core i7 y 16 GB de RAM para el desarrollo local y pruebas iniciales.
- Una (1) GPU Tesla T4 (14.6 GB VRAM) proporcionada por la plataforma Kaggle para el entrenamiento del modelo.

Software y tecnologías:
- Sistema operativo Windows 11 para el desarrollo local.
- Python 3.12 como lenguaje de programación principal.
- PyTorch 2.x y Torchvision para la construcción y entrenamiento del modelo.
- Bibliotecas auxiliares: scikit-learn, NumPy, Pandas, Matplotlib, Seaborn, OpenCV, Streamlit.
- Dataset público "140k Real and Fake Faces" de Kaggle para entrenamiento, validación y prueba.
- Plataforma Kaggle (con GPU) para el entrenamiento del modelo."""

    for i, p in enumerate(paragraphs):
        if "Recursos necesarios" in (p.text or ""):
            j = i + 1
            # Saltar párrafos de contenido antiguo y reemplazar
            while j < len(paragraphs):
                pj = paragraphs[j]
                if "Localización" in (pj.text or ""):
                    break
                if pj.text and pj.text.strip():
                    for run in pj.runs:
                        run.text = ''
                    if pj.runs:
                        if not pj.runs[0].text:
                            pj.runs[0].text = recursos_nuevo
                    else:
                        pj.add_run(recursos_nuevo)
                    break
                j += 1
            break
    
    # CORREGIR sección Localización
    localizacion_nueva = """La fase de experimentación y entrenamiento del modelo se ejecutó en la plataforma Kaggle, haciendo uso de su infraestructura cloud con aceleración GPU (Tesla T4). El desarrollo, las pruebas locales y la implementación del prototipo web se realizaron en el entorno de desarrollo local del investigador. El prototipo final se despliega como una aplicación web accesible a través del navegador, sin requerir infraestructura especializada por parte del usuario final."""
    
    for i, p in enumerate(paragraphs):
        if "Localización del proyecto" in (p.text or ""):
            j = i + 1
            while j < len(paragraphs):
                pj = paragraphs[j]
                if "Cronograma" in (pj.text or ""):
                    break
                if pj.text and pj.text.strip():
                    for run in pj.runs:
                        run.text = ''
                    if pj.runs:
                        pj.runs[0].text = localizacion_nueva
                    else:
                        pj.add_run(localizacion_nueva)
                    break
                j += 1
            break
    
    # Guardar documento mejorado
    doc.save(dest_path)
    print(f"Documento mejorado guardado en: {dest_path}")
    return dest_path

if __name__ == "__main__":
    path = create_improved_document()
    print(f"\nArchivo creado: {path}")
