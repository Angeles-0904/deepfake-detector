"""
Script v3: Mejora la Metodologia con vinculacion objetivo-metodologia,
limpia contenido residual de IoT/Raspberry Pi, y actualiza las Referencias.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ============================================================
# CONTENIDO MEJORADO
# ============================================================

METODOLOGIA_OBJETIVOS = """
La presente investigacion empleo una metodologia de tipo aplicada con enfoque cuantitativo y nivel explicativo-predictivo y tecnologico, desarrollada bajo un diseno experimental computacional. A continuacion se describe la relacion entre cada objetivo especifico y la metodologia empleada para su cumplimiento:

Objetivo Especifico 1: "Analizar las caracteristicas presentes en imagenes reales y deepfakes mediante tecnicas de procesamiento digital de imagenes para identificar patrones relevantes para su clasificacion."

Para el cumplimiento de este objetivo se aplico la metodologia de analisis exploratorio de datos (EDA por sus siglas en ingles). Esta metodologia, descrita por Tukey (1977) y adaptada por Hernandez-Sampieri y Mendoza (2018) como parte del enfoque cuantitativo, permitio examinar visual y estadisticamente las caracteristicas de las imagenes del dataset "140k Real and Fake Faces". Se analizaron distribuciones de intensidad de pixeles, histogramas de color en los canales RGB, niveles de ruido, frecuencia de texturas y artefactos de compresion. Se utilizaron las bibliotecas OpenCV y Matplotlib en Python para el procesamiento y visualizacion de las imagenes. Este analisis exploratorio sento las bases para disenar el preprocesamiento que se aplicaria en etapas posteriores. Según Hernandez-Sampieri y Mendoza (2018), el analisis exploratorio es la primera etapa en investigaciones cuantitativas que buscan describir las caracteristicas de los datos antes de someterlos a pruebas mas complejas.

Objetivo Especifico 2: "Implementar una red neuronal convolucional basada en Transfer Learning para la clasificacion automatica de imagenes reales y manipuladas digitalmente."

Para este objetivo se empleo la metodologia de Transfer Learning (aprendizaje por transferencia), que consiste en reutilizar los pesos de un modelo preentrenado en una tarea base y adaptarlos a una tarea objetivo relacionada. Esta metodologia, formalizada por Pan y Yang (2010) y ampliamente adoptada en vision por computadora, permitio tomar la arquitectura DenseNet-121 preentrenada en ImageNet (con 1.2 millones de imagenes y 1000 clases) y adaptarla a la clasificacion binaria (real vs. deepfake) mediante fine-tuning. Se congelaron las capas convolucionales inferiores para preservar las caracteristicas generales (bordes, texturas, formas) y se reemplazo el clasificador por una nueva cabeza fully connected. La implementacion se realizo con PyTorch y Torchvision. Este enfoque es consistente con lo senalado por Goodfellow, Bengio y Courville (2016), quienes describen el transfer learning como una estrategia efectiva cuando los datos de la tarea objetivo son insuficientes para entrenar una red desde cero.

Objetivo Especifico 3: "Evaluar el desempeno del modelo mediante metricas de precision, recall, F1-Score y AUC para determinar su efectividad en la deteccion de deepfakes."

Se aplico la metodologia de evaluacion de clasificadoresbinarios mediante metricas estandar de rendimiento, conforme a los principios establecidos por Sokolova y Lapalme (2009) y Flach (2012). Esta metodologia implico: (a) calculo de la matriz de confusion (verdaderos positivos, verdaderos negativos, falsos positivos, falsos negativos); (b) derivacion de metricas derivadas: accuracy, precision, recall (sensibilidad), especificidad, F1-score; (c) generacion de la curva ROC y calculo del AUC (Area Under the Curve). La evaluacion se realizo sobre un conjunto de prueba independiente de 20,000 imagenes que el modelo no habia visto durante el entrenamiento. Segun Hernandez-Sampieri y Mendoza (2018), la evaluacion mediante metricas cuantitativas constituye un procedimiento central en investigaciones de enfoque cuantitativo y nivel explicativo-predictivo.

Objetivo Especifico 4: "Aplicar tecnicas de Inteligencia Artificial Explicable mediante Grad-CAM para generar explicaciones visuales e interpretables sobre las decisiones tomadas por el modelo de deteccion de deepfakes."

Para este objetivo se empleo la metodologia de Inteligencia Artificial Explicable (XAI, por sus siglas en ingles), especificamente la tecnica Grad-CAM (Gradient-weighted Class Activation Mapping) desarrollada por Selvaraju et al. (2017). La metodologia XAI busca hacer interpretables las decisiones de modelos de caja negra como las redes neuronales profundas. El procedimiento consistio en: (a) propagar una imagen hacia adelante a traves del modelo para obtener la prediccion; (b) calcular los gradientes de la clase predecida con respecto a los mapas de activacion de la ultima capa convolucional; (c) ponderar los mapas de activacion por los gradientes promedio; y (d) superponer el mapa de calor resultante sobre la imagen original. Segun Arrieta et al. (2020), la XAI se ha convertido en un area prioritaria de investigacion debido a la necesidad de confianza y transparencia en sistemas de inteligencia artificial.

Objetivo Especifico 5: "Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificacion de contenido digital manipulado."

Se aplico la metodologia de validacion funcional y de usabilidad, basada en dos componentes. El primer componente fue la validacion funcional, que consistio en verificar que el prototipo web (implementado en Streamlit) recibiera correctamente la imagen de entrada, ejecutara el modelo de clasificacion, generara los mapas Grad-CAM y mostrara los resultados de forma clara e interactiva. El segundo componente fue la evaluacion de usabilidad, aplicando el cuestionario System Usability Scale (SUS) desarrollado por Brooke (1996), que consta de 10 preguntas con puntuacion de 1 a 5. La muestra para esta evaluacion fue no probabilisticade tipo intencional o por conveniencia, seleccionando entre 15 y 30 usuarios con conocimientos basicos en tecnologia. Segun Ñaupas et al. (2018), la validacion funcional y de usabilidad constituye un procedimiento propio de investigaciones de tipo aplicado y nivel tecnologico.
"""

REFERENCIAS_TEXT = """
Arrieta, A. B., Diaz-Rodriguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., ... & Herrera, F. (2020). Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 58, 82-115.

Brooke, J. (1996). SUS: A quick and dirty usability scale. Usability Evaluation in Industry, 189(194), 4-7.

Flach, P. (2012). Machine Learning: The Art and Science of Algorithms that Make Sense of Data. Cambridge University Press.

Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.

Hernandez-Sampieri, R. y Mendoza, C. (2018). Metodologia de la investigacion: Las rutas cuantitativa, cualitativa y mixta. McGraw-Hill Interamericana.

Naupas, H., Valdivia, M., Palacios, J., y Romero, H. (2018). Metodologia de la investigacion cuantitativa-cualitativa y redaccion de la tesis (5a ed.). Ediciones de la U.

Pan, S. J., & Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering, 22(10), 1345-1359.

Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. Proceedings of the IEEE International Conference on Computer Vision (ICCV), 618-626.

Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427-437.

Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley.
"""


def find_paragraph_after_heading(doc, heading_text, start_idx=0):
    """Find a heading paragraph and return the next content paragraph index."""
    for i, p in enumerate(doc.paragraphs):
        if i < start_idx:
            continue
        if heading_text.lower() in (p.text or '').lower():
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    return j
            return None
    return None


def clear_and_set_text(paragraph, text):
    """Clear all runs in a paragraph and set new text."""
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def clear_paragraph_range(doc, start_idx, end_idx):
    """Clear all paragraphs in a range (inclusive start, exclusive end)."""
    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = ''
        else:
            p.add_run('')


def create_improved_document():
    source_path = "C:/Users/SHAMELY/Documents/Downloads/samely.docx"
    dest_path = "D:/IX/detectorIA/documentacion/PROYECTO_INVESTIGACION_MEJORADO.docx"

    doc = Document(source_path)

    print("=== MEJORANDO DOCUMENTO v3 ===")

    # ===================================================================
    # 1. MEJORAR METODOLOGIA - Agregar seccion de metodologia por objetivos
    # ===================================================================
    # Buscar la seccion "Metodologia de investigacion" o "Tipo y enfoque"
    # Insertar la explicacion de metodologia por objetivos DESPUES de "Tipo y enfoque de investigacion"
    # pero antes de "Diseno de investigacion"

    tipo_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Tipo y enfoque de investigacion" in (p.text or '') and i > 80:
            tipo_idx = i
            break

    if tipo_idx:
        print(f"Seccion 'Tipo y enfoque' encontrada en indice {tipo_idx}")
        # Insertar la explicacion de metodologia por objetivos justo antes de "Diseno de investigacion"
        disenio_idx = None
        for j in range(tipo_idx + 1, len(doc.paragraphs)):
            if "Diseno de investigacion" in (doc.paragraphs[j].text or ''):
                disenio_idx = j
                break

        if disenio_idx:
            print(f"Seccion 'Diseno de investigacion' encontrada en indice {disenio_idx}")
            # Insertar la metodologia por objetivos ANTES de "Diseno de investigacion"
            # Crear parrafos con el contenido
            ref_para = doc.paragraphs[tipo_idx + 1]
            lines = METODOLOGIA_OBJETIVOS.strip().split('\n')
            
            # Insertar parrafos de atras hacia adelante para mantener indices
            for line in reversed(lines):
                line = line.strip()
                if line:
                    new_p = deepcopy(ref_para._element)
                    # Remove existing runs
                    for r in list(new_p.findall(qn('w:r'))):
                        new_p.remove(r)
                    # Insert before "Diseno de investigacion" paragraph
                    doc.paragraphs[disenio_idx]._element.addprevious(new_p)
                    # Agregar run <w:r> con <w:t> al nuevo elemento
                    r_element = OxmlElement('w:r')
                    t_element = OxmlElement('w:t')
                    t_element.set(qn('xml:space'), 'preserve')
                    t_element.text = line
                    r_element.append(t_element)
                    new_p.append(r_element)

    # ===================================================================
    # 2. LIMPIAR SECCION RECURSOS (eliminar IoT/Raspberry Pi)
    # ===================================================================
    print("\nLimpiando seccion Recursos...")
    recursos_start = None
    for i, p in enumerate(doc.paragraphs):
        if "Recursos necesarios" in (p.text or ''):
            recursos_start = i
            break

    if recursos_start:
        print(f"Seccion 'Recursos necesarios' encontrada en indice {recursos_start}")
        # Encontrar donde termina (proxima seccion principal)
        recursos_end = None
        for j in range(recursos_start + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text or ''
            if text.strip() and ('Cuerpo' in doc.paragraphs[j].style.name or 'Heading' in doc.paragraphs[j].style.name):
                recursos_end = j
                break
        if not recursos_end:
            recursos_end = len(doc.paragraphs)
        
        print(f"Rango de recursos: {recursos_start + 1} a {recursos_end}")
        
        # Limpiar todos los parrafos de recursos
        recursos_content = """Para la ejecucion del presente proyecto de investigacion se requirieron los siguientes recursos, organizados en equipos de computo, software y datos.

Equipos:
- Una (1) laptop con procesador Intel Core i7 y 16 GB de RAM para el desarrollo local y pruebas iniciales.
- Una (1) GPU Tesla T4 (14.6 GB VRAM) proporcionada por la plataforma Kaggle para el entrenamiento del modelo.

Software y tecnologias:
- Sistema operativo Windows 11 para el desarrollo local.
- Python 3.12 como lenguaje de programacion principal.
- PyTorch 2.x y Torchvision para la construccion y entrenamiento del modelo.
- Bibliotecas auxiliares: scikit-learn, NumPy, Pandas, Matplotlib, Seaborn, OpenCV, Streamlit.
- Dataset publico "140k Real and Fake Faces" de Kaggle para entrenamiento, validacion y prueba.
- Plataforma Kaggle (con GPU) para el entrenamiento del modelo.

Infraestructura:
- Espacio de trabajo con computadora y conexion a internet de banda ancha.
- Cuenta en Kaggle para acceso a datasets y GPU."""
        
        # Reemplazar el primer parrafo de contenido despues del titulo
        first_content_idx = None
        for j in range(recursos_start + 1, recursos_end):
            if doc.paragraphs[j].text.strip():
                first_content_idx = j
                break
        
        if first_content_idx:
            clear_and_set_text(doc.paragraphs[first_content_idx], recursos_content)
            # Limpiar los demas parrafos
            for k in range(first_content_idx + 1, recursos_end):
                doc.paragraphs[k].text = ''
                for run in doc.paragraphs[k].runs:
                    run.text = ''

    # ===================================================================
    # 3. LIMPIAR SECCION LOCALIZACION (residual)
    # ===================================================================
    print("Actualizando seccion Localizacion...")
    localizacion_content = """La fase de experimentacion y entrenamiento del modelo se ejecuto en la plataforma Kaggle, haciendo uso de su infraestructura cloud con aceleracion GPU (Tesla T4). El desarrollo, las pruebas locales y la implementacion del prototipo web se realizaron en el entorno de desarrollo local del investigador. El prototipo final se despliega como una aplicacion web accesible a traves del navegador, sin requerir infraestructura especializada por parte del usuario final."""

    idx = find_paragraph_after_heading(doc, "Localizacion del proyecto")
    if idx:
        clear_and_set_text(doc.paragraphs[idx], localizacion_content)

    # ===================================================================
    # 4. AGREGAR REFERENCIAS BIBLIOGRAFICAS
    # ===================================================================
    print("Actualizando seccion Referencias...")
    
    # Buscar la seccion de Referencias
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Referencias" in (p.text or '') and i > 100:
            style_name = p.style.name
            if 'Cuerpo' in style_name or 'Heading' in style_name:
                ref_idx = i
                break

    if ref_idx is None:
        # Buscar al final del documento
        for i in range(len(doc.paragraphs) - 1, 0, -1):
            if "Referencias" in (doc.paragraphs[i].text or ''):
                ref_idx = i
                break

    if ref_idx:
        print(f"Seccion 'Referencias' encontrada en indice {ref_idx}")
        # Limpiar parrafos de referencias existentes despues del titulo
        ref_end = None
        for j in range(ref_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text or ''
            if text.strip() and ('Cuerpo' in doc.paragraphs[j].style.name or 'Heading' in doc.paragraphs[j].style.name):
                ref_end = j
                break
        if not ref_end:
            ref_end = len(doc.paragraphs)
        
        # Eliminar contenido de referencias existente
        for j in range(ref_idx + 1, ref_end):
            doc.paragraphs[j].text = ''
            for run in doc.paragraphs[j].runs:
                run.text = ''
        
        # Agregar nuevas referencias, cada una como un parrafo separado
        refs = [
            "Arrieta, A. B., Diaz-Rodriguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., ... & Herrera, F. (2020). Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 58, 82-115.",
            "Brooke, J. (1996). SUS: A quick and dirty usability scale. Usability Evaluation in Industry, 189(194), 4-7.",
            "Flach, P. (2012). Machine Learning: The Art and Science of Algorithms that Make Sense of Data. Cambridge University Press.",
            "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
            "Hernandez-Sampieri, R. y Mendoza, C. (2018). Metodologia de la investigacion: Las rutas cuantitativa, cualitativa y mixta. McGraw-Hill Interamericana.",
            "Naupas, H., Valdivia, M., Palacios, J., y Romero, H. (2018). Metodologia de la investigacion cuantitativa-cualitativa y redaccion de la tesis (5a ed.). Ediciones de la U.",
            "Pan, S. J., & Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering, 22(10), 1345-1359.",
            "Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. Proceedings of the IEEE International Conference on Computer Vision (ICCV), 618-626.",
            "Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427-437.",
            "Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley."
        ]
        
        # Obtener el elemento del titulo de Referencias
        ref_title_element = doc.paragraphs[ref_idx]._element
        
        # Insertar referencias despues del titulo usando XML directo
        for ref_text in refs:
            # Crear elemento <w:p>
            p_element = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'Normal')
            pPr.append(pStyle)
            p_element.append(pPr)
            
            # Crear <w:r>
            r_element = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            r_element.append(rPr)
            
            # Crear <w:t> con el texto
            t_element = OxmlElement('w:t')
            t_element.set(qn('xml:space'), 'preserve')
            t_element.text = ref_text
            r_element.append(t_element)
            p_element.append(r_element)
            
            # Insertar despues del titulo de Referencias
            ref_title_element.addnext(p_element)

    # ===================================================================
    # 5. BARRIDO FINAL: Eliminar cualquier texto residual de IoT/Raspberry Pi
    # ===================================================================
    print("\nBarrido final de contenido residual...")
    residual_keywords = [
        "Raspberry Pi", "raspberry", "IoT", "BoT-IoT", "TensorFlow Lite", 
        "TFLite", "TensorFlow 2.x", "Wireshark", "tcpdump",
        "microSD", "camara IP", "sensor de temperatura",
        "Raspberry Pi OS", "dispositivo IoT", "mesa de experimentacion"
    ]
    
    cambios = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text or ''
        for kw in residual_keywords:
            if kw.lower() in text.lower():
                # Limpiar este parrafo
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = ''
                else:
                    p.add_run('')
                cambios += 1
                print(f"  Limpiado [{i}]: contenido con '{kw}'")
                break

    if cambios == 0:
        print("  No se encontro contenido residual de IoT/Raspberry Pi.")

    # ===================================================================
    # 6. Guardar documento
    # ===================================================================
    doc.save(dest_path)
    print(f"\nDocumento guardado en: {dest_path}")
    print(f"Cambios realizados: {cambios} parrafos residuales limpiados")
    
    # Verificacion final
    print("\n=== VERIFICACION FINAL ===")
    verificar_keywords = ["Raspberry", "IoT", "BoT-IoT", "TensorFlow Lite", "Wireshark"]
    for i, p in enumerate(doc.paragraphs):
        for kw in verificar_keywords:
            if kw.lower() in (p.text or '').lower():
                print(f"  ADVERTENCIA: [{i}] AUN contiene '{kw}': {(p.text or '')[:100]}...")

    return dest_path


if __name__ == "__main__":
    path = create_improved_document()
    print(f"\nArchivo creado: {path}")
