"""
Script v5 - Version COMPLETA y DEFINITIVA
Combina todas las mejoras desde el original samely.docx:
1. 15 Antecedentes Internacionales (v2)
2. Metodologia base mejorada (v2)
3. Metodologia por Objetivos (v4 - orden correcto)
4. Impactos corregidos (v2)
5. Localizacion corregida (v2)
6. Recursos corregidos (v4)
7. Referencias actualizadas (v4)
8. Barrido final
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import unicodedata


def normalize(text):
    if text is None:
        return ''
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii').lower().strip()


def create_new_paragraph_element(text, style='Normal'):
    """Create a <w:p> element with <w:r><w:t> containing the text."""
    p_element = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    p_element.append(pPr)
    r_element = OxmlElement('w:r')
    t_element = OxmlElement('w:t')
    t_element.set(qn('xml:space'), 'preserve')
    t_element.text = text
    r_element.append(t_element)
    p_element.append(r_element)
    return p_element


def insert_after(element, texts, style='Normal'):
    """Insert paragraphs after element, preserving order."""
    for text in reversed(texts):
        p_element = create_new_paragraph_element(text, style)
        element.addnext(p_element)


def clear_range(doc, start, end):
    """Remove paragraph elements from XML in range [start, end)."""
    for i in range(min(end, len(doc.paragraphs)) - 1, start - 1, -1):
        if i < len(doc.paragraphs):
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)


def find_first(doc, text, after=0):
    """Find first paragraph index containing text."""
    nt = normalize(text)
    for i, p in enumerate(doc.paragraphs):
        if i < after:
            continue
        if nt in normalize(p.text or ''):
            return i
    return None


# ============================================================
# CONTENIDO
# ============================================================

ANTECEDENTES = [
    "Abdul-Hafiz & Sari (2025) desarrollaron un framework robusto y explicable para detectar deepfakes en imagenes utilizando transfer learning con mecanismos de atencion (CBAM). Emplearon 5 redes neuronales convolucionales en paralelo, las combinaron con XGBoost y aplicaron Grad-CAM para generar mapas de calor. Alcanzaron una precision del 97.25% en un dataset de 140,000 imagenes.",
    "Rakesh Kumar (2025) propuso un sistema de deteccion de deepfakes basado en redes neuronales convolucionales con mecanismos de atencion, integrando tecnicas de XAI como Grad-CAM, LIME y SHAP. Comparo multiples arquitecturas pre-entrenadas (EfficientNet, ResNet y VGG) y alcanzo una precision del 94.7% y un AUC de 0.967.",
    "Mohit Kumar (2025) desarrollo un detector de deepfakes centrado en la robustez del modelo, evaluando bajo condiciones adversas como compresion JPEG, ruido y ataques adversariales usando EfficientNet, ResNet y VGG con Grad-CAM, LIME y SHAP. Mantubo una alta precision (94.7%) incluso con degradaciones.",
    "Aribe Jr. (2025) propuso un enfoque hibrido que fusiona caracteristicas forenses tradicionales con aprendizaje profundo, integrando analisis de ruido PRNU, artefactos JPEG y coeficientes DCT con ResNet-50 y Vision Transformer, usando Grad-CAM para mapas de calor forenses. Logro un overlap del 82% entre mapas de calor y regiones manipuladas.",
    "Lipianina-Honcharenko et al. (2025) implementaron un metodo de conjunto basado en cinco CNNs (ResNet50, EfficientNetB0, Xception, InceptionV3 y FaceNet) con seleccion de golden frames, combinadas mediante XGBoost, logrando una precision del 91.14%.",
    "Raikwar et al. (2025) desarrollaron un modelo hibrido CNN-transformer para deteccion de deepfakes, obteniendo una precision del 94.8% y un F1-score del 93.9%.",
    "Singh & Kumar (2025) propusieron una arquitectura CNN-RNN para deteccion de deepfakes, combinando extraccion de caracteristicas espaciales con modelado de dependencias temporales.",
    "Wang et al. (2020) desarrollaron un detector universal de imagenes generadas por CNNs, evaluando en 11 generadores (ProGAN, StyleGAN, BigGAN), demostrando deteccion independiente de la arquitectura generadora.",
    "Selvaraju et al. (2017) introdujeron Grad-CAM, tecnica de explicabilidad visual que genera mapas de calor desde los gradientes de la ultima capa convolucional de una CNN. Estandar de facto para explicabilidad en CNNs.",
    "Rossler et al. (2019) introdujeron FaceForensics++, dataset de referencia con 1,000 videos reales y 4,000 manipulados con 4 tecnicas (Deepfakes, Face2Face, FaceSwap, NeuralTextures), en 3 niveles de compresion.",
    "Li et al. (2020) introdujeron Celeb-DF v2, dataset de deepfakes de alta calidad con 590 videos reales y 5,639 deepfake generados con intercambio facial avanzado.",
    "Tolosana et al. (2020) publicaron una revision exhaustiva del estado del arte en deteccion de deepfakes y manipulacion facial, cubriendo generacion, deteccion, datasets y desafios.",
    "Guarnera, Giudice y Battiato (2020) desarrollaron DeepFake Detection by Analyzing Convolutional Traces, identificando huellas forenses ocultas en procesos convolucionales de generacion de rostros sinteticos.",
    "De Lima et al. (2020) realizaron Deepfake Detection Using Spatiotemporal Convolutional Networks, mejorando la deteccion mediante analisis conjunto de informacion espacial y temporal.",
    "Tariq et al. (2025) desarrollaron From Prediction to Explanation: framework multimodal, explicable e interactivo para deteccion de deepfakes orientado a usuarios no especializados."
]

METODOLOGIA_BASE = {
    "Tipo y enfoque": "La investigacion fue de enfoque cuantitativo, empleando datos numericos, metricas de desempeno y procedimientos estadisticos. Hernandez-Sampieri y Mendoza (2018) senalan que el enfoque cuantitativo permite medir variables, contrastar hipotesis y analizar resultados mediante procedimientos estadisticos. Por su finalidad, fue de tipo aplicado, usando conocimientos de vision por computadora, CNN, transfer learning y XAI para desarrollar una solucion tecnologica. Segun Naupas et al. (2018), la investigacion aplicada busca resolver necesidades especificas. El nivel fue explicativo-predictivo y tecnologico: explico las caracteristicas visuales que diferencian imagenes autenticas de manipuladas, predijo la probabilidad de clase, y culmino en un prototipo funcional con Streamlit.",
    "Diseno": "Se utilizo un diseno experimental computacional, modificando condiciones del entrenamiento (arquitectura, tasa de aprendizaje, epocas, congelamiento de capas, aumento de datos) para observar su efecto en el rendimiento. El procedimiento se organizo en seis fases: (1) seleccion y preparacion del dataset, (2) preprocesamiento, (3) entrenamiento de DenseNet-121 con transfer learning, (4) evaluacion mediante metricas, (5) explicaciones con Grad-CAM, (6) implementacion del prototipo web.",
    "Variables": "La variable independiente fue la configuracion del sistema basado en DenseNet-121 con transfer learning y Grad-CAM. La variable dependiente fue el desempeno en la deteccion explicable de deepfakes, medido mediante accuracy, precision, recall, F1-score, AUC-ROC y tiempo de inferencia.",
    "Poblacion": "La poblacion fueron 140,000 imagenes faciales del dataset '140k Real and Fake Faces' (Kaggle), balanceadas: 70,000 reales y 70,000 deepfake. Se uso muestra censal (totalidad de imagenes). Division estratificada: 100,000 entrenamiento (71.4%), 20,000 validacion (14.3%), 20,000 prueba (14.3%). Se evito fuga de informacion entre conjuntos.",
    "Tecnicas": "La tecnica principal fue la observacion computacional, complementada con analisis experimental y pruebas funcionales. Instrumentos: ficha de registro del dataset, reporte de metricas, matriz de confusion, curvas ROC, mapas Grad-CAM y registro de tiempos. Se uso PyTorch 2.x, Torchvision, Scikit-learn, Matplotlib, Seaborn, OpenCV.",
    "Procedimiento": "El procedimiento se ejecuto en Kaggle con GPU Tesla T4 (14.6 GB VRAM). Etapas: (1) Descarga del dataset '140k Real and Fake Faces', organizacion en carpetas por clase y conjunto. (2) Preprocesamiento: redimensionamiento a 224x224, normalizacion con media/desviacion de ImageNet, aumento de datos (rotaciones +/-15 grados, volteo horizontal). (3) Arquitectura: DenseNet-121 preentrenada en ImageNet, clasificador reemplazado por FC 1024->512 ReLU + dropout 0.3 + FC 512->2. Total: 7,479,682 parametros, 1,011,778 entrenables (13.5%). (4) Entrenamiento: optimizador Adam, tasa 1e-4 (clasificador) y 1e-5 (fine-tuning), CrossEntropyLoss, ReduceLROnPlateau, 14 epocas, early stopping con paciencia 7. (5) Evaluacion: 97.06% accuracy, 0.9966 AUC, F1 macro 0.97 en 20,000 imagenes de prueba. (6) Implementacion: app web Streamlit con clasificacion, mapas Grad-CAM, explicaciones textuales y pruebas de robustez.",
    "Analisis": "Metricas en prueba: Accuracy 97.06%, Precision macro 0.97, Recall macro 0.97, F1-score macro 0.97, AUC-ROC 0.9966. Matriz de confusion analizada. Mapas Grad-CAM generados para visualizar regiones influyentes. Umbral satisfactorio: F1 y AUC >= 0.85.",
    "Etica": "Se usaron datasets publicos con fines academicos, respetando licencias. No se emplearon imagenes privadas sin consentimiento. Resultados presentados de forma transparente. Sistema descrito como herramienta de apoyo, no sustituto de peritaje digital."
}

METODOLOGIA_OBJETIVOS = """La presente investigacion empleo una metodologia de tipo aplicada con enfoque cuantitativo y nivel explicativo-predictivo y tecnologico, desarrollada bajo un diseno experimental computacional. A continuacion se describe la relacion entre cada objetivo especifico y la metodologia empleada para su cumplimiento:

Objetivo Especifico 1: 'Analizar las caracteristicas presentes en imagenes reales y deepfakes mediante tecnicas de procesamiento digital de imagenes para identificar patrones relevantes para su clasificacion.'

Para este objetivo se aplico la metodologia de Analisis Exploratorio de Datos (EDA), descrita por Tukey (1977) y adaptada por Hernandez-Sampieri y Mendoza (2018). Se examinaron las imagenes del dataset '140k Real and Fake Faces' analizando distribuciones de intensidad de pixeles, histogramas RGB, niveles de ruido, texturas y artefactos de compresion con OpenCV y Matplotlib.

Objetivo Especifico 2: 'Implementar una red neuronal convolucional basada en Transfer Learning para la clasificacion automatica de imagenes reales y manipuladas digitalmente.'

Se empleo la metodologia de Transfer Learning, formalizada por Pan y Yang (2010). Se tomo DenseNet-121 preentrenada en ImageNet, se congelaron las capas inferiores y se reemplazo el clasificador por una cabeza fully connected. Implementacion con PyTorch y Torchvision. Goodfellow, Bengio y Courville (2016) describen el transfer learning como efectivo cuando los datos de la tarea objetivo son insuficientes para entrenar desde cero.

Objetivo Especifico 3: 'Evaluar el desempeno del modelo mediante metricas de precision, recall, F1-Score y AUC para determinar su efectividad en la deteccion de deepfakes.'

Se aplico la metodologia de evaluacion de clasificadores binarios (Sokolova y Lapalme, 2009; Flach, 2012). Se calculo la matriz de confusion, accuracy, precision, recall, especificidad, F1-score, curva ROC y AUC. La evaluacion se realizo sobre 20,000 imagenes de prueba independientes.

Objetivo Especifico 4: 'Aplicar tecnicas de Inteligencia Artificial Explicable mediante Grad-CAM para generar explicaciones visuales e interpretables sobre las decisiones tomadas por el modelo.'

Se empleo la metodologia XAI mediante Grad-CAM (Selvaraju et al., 2017). El procedimiento: propagacion hacia adelante, calculo de gradientes de la clase respecto a mapas de activacion de la ultima capa convolucional, ponderacion y superposicion del mapa de calor. Arrieta et al. (2020) destacan la XAI como area prioritaria para confianza y transparencia en IA.

Objetivo Especifico 5: 'Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificacion de contenido digital manipulado.'

Se aplico validacion funcional y de usabilidad. La validacion funcional verifico que el prototipo Streamlit recibiera imagenes, ejecutara el modelo, generara mapas Grad-CAM y mostrara resultados. La usabilidad se evaluo mediante System Usability Scale (SUS) de Brooke (1996). Muestra no probabilistica intencional de 15 a 30 usuarios. Naupas et al. (2018) senalan que la validacion funcional y de usabilidad es propia de investigaciones aplicadas de nivel tecnologico."""

IMPACTOS = {
    "ciencia": "El proyecto genero impactos en ciencia y tecnologia al aportar evidencia sobre la aplicacion de CNN con transfer learning y XAI en deteccion de deepfakes. En lo cientifico, amplio el conocimiento sobre patrones visuales asociados a contenidos manipulados y el comportamiento de DenseNet-121. En lo tecnologico, se desarrollo un prototipo web (Streamlit) para clasificar imagenes y visualizar decisiones mediante Grad-CAM, usando tecnologias abiertas (Python, PyTorch, OpenCV) que favorecen la reproducibilidad.",
    "economicos": "Se demuestra que es viable desarrollar un detector de deepfakes funcional con hardware accesible. El modelo se ejecuta en CPUs convencionales con tiempo de inferencia razonable, sin requerir GPUs especializadas, reduciendo la barrera de entrada para instituciones educativas y pequenas organizaciones.",
    "sociales": "Los deepfakes pueden difundir desinformacion, suplantar identidades y erosionar la confianza en evidencias audiovisuales. Esta investigacion contribuye a mitigar estos riesgos proporcionando una herramienta accesible para identificar contenido manipulado, fortaleciendo la alfabetizacion digital.",
    "ambientales": "El impacto ambiental se mitigo usando la plataforma Kaggle con hardware eficiente, reduciendo epocas mediante early stopping y aplicando transfer learning para disminuir el tiempo de entrenamiento. El prototipo final se ejecuta en hardware de consumo general."
}

LOCALIZACION = "La fase de experimentacion y entrenamiento se ejecuto en la plataforma Kaggle con GPU Tesla T4. El desarrollo local y la implementacion del prototipo web se realizaron en el entorno del investigador. El prototipo final es una aplicacion web accesible via navegador, sin requerir infraestructura especializada."

RECURSOS = [
    "Para la ejecucion del presente proyecto se requirieron los siguientes recursos:",
    "Equipos:",
    "- Una (1) laptop con Intel Core i7 y 16 GB de RAM para desarrollo local.",
    "- Una (1) GPU Tesla T4 (14.6 GB VRAM) en Kaggle para entrenamiento.",
    "Software:",
    "- Windows 11 para desarrollo local.",
    "- Python 3.12, PyTorch 2.x, Torchvision, Scikit-learn, NumPy, Pandas.",
    "- Matplotlib, Seaborn, OpenCV, Streamlit.",
    "- Dataset '140k Real and Fake Faces' de Kaggle.",
    "- Plataforma Kaggle con GPU Tesla T4.",
    "Infraestructura:",
    "- Espacio de trabajo con computadora y conexion a internet.",
    "- Cuenta en Kaggle para acceso a datasets y GPU."
]

REFERENCIAS = [
    "Arrieta, A. B., Diaz-Rodriguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., ... & Herrera, F. (2020). Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 58, 82-115.",
    "Brooke, J. (1996). SUS: A quick and dirty usability scale. Usability Evaluation in Industry, 189(194), 4-7.",
    "Flach, P. (2012). Machine Learning: The Art and Science of Algorithms that Make Sense of Data. Cambridge University Press.",
    "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
    "Hernandez-Sampieri, R. y Mendoza, C. (2018). Metodologia de la investigacion: Las rutas cuantitativa, cualitativa y mixta. McGraw-Hill Interamericana.",
    "Naupas, H., Valdivia, M., Palacios, J., y Romero, H. (2018). Metodologia de la investigacion cuantitativa-cualitativa y redaccion de la tesis (5a ed.). Ediciones de la U.",
    "Pan, S. J., & Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering, 22(10), 1345-1359.",
    "Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. Proceedings of the IEEE International Conference on Computer Vision (ICCV), 618-626.",
    "Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427-437.",
    "Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley."
]

RESIDUAL_KEYWORDS = [
    "Raspberry Pi", "raspberry", "IoT", "BoT-IoT", "TensorFlow Lite",
    "TFLite", "TensorFlow 2.x", "Wireshark", "tcpdump",
    "microSD", "camara IP", "sensor de temperatura",
    "Raspberry Pi OS", "dispositivo IoT", "mesa de experimentacion",
    "TabTransformer", "phishing", "ciberseguridad", "bot-iot", "malicios"
]


def create_document():
    source = "C:/Users/SHAMELY/Documents/Downloads/samely.docx"
    dest = "D:/IX/detectorIA/documentacion/PROYECTO_INVESTIGACION_MEJORADO.docx"

    doc = Document(source)
    total_changes = 0

    print("=== CREANDO DOCUMENTO V5 (COMPLETO) ===\n")

    # ===================================================================
    # 1. ANTECEDENTES INTERNACIONALES
    # ===================================================================
    print("[1/7] Antecedentes Internacionales...")
    ant_idx = find_first(doc, "Antecedentes internacionales")
    if ant_idx is not None:
        nac_idx = find_first(doc, "Antecedentes nacionales", ant_idx + 1)
        if nac_idx is None:
            nac_idx = ant_idx + 1 + 30
        
        # Limpiar antecedentes existentes
        content_paras = []
        for j in range(ant_idx + 1, nac_idx):
            if doc.paragraphs[j].text.strip():
                content_paras.append(j)
        
        # Reemplazar primeros N parrafos
        for idx, para_idx in enumerate(content_paras):
            if idx < len(ANTECEDENTES):
                p = doc.paragraphs[para_idx]
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = ANTECEDENTES[idx]
                else:
                    p.add_run(ANTECEDENTES[idx])
        
        # Si hay mas antecedentes que parrafos existentes, agregar
        if len(ANTECEDENTES) > len(content_paras):
            last_element = doc.paragraphs[nac_idx - 1]._element
            for extra_idx in range(len(content_paras), len(ANTECEDENTES)):
                p_element = create_new_paragraph_element(ANTECEDENTES[extra_idx], 'Normal')
                last_element.addnext(p_element)
                last_element = p_element
        
        total_changes += len(ANTECEDENTES)
        print(f"  Insertados {len(ANTECEDENTES)} antecedentes internacionales")
    else:
        print("  ERROR: No se encontro 'Antecedentes internacionales'")

    # ===================================================================
    # 2. METODOLOGIA BASE (v2)
    # ===================================================================
    print("\n[2/7] Metodologia base...")
    for title, content in METODOLOGIA_BASE.items():
        idx = find_first(doc, title)
        if idx is not None:
            for j in range(idx + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    p = doc.paragraphs[j]
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = content
                    else:
                        p.add_run(content)
                    total_changes += 1
                    print(f"  '{title}' actualizado")
                    break
        else:
            print(f"  '{title}' no encontrado")

    # ===================================================================
    # 3. METODOLOGIA POR OBJETIVOS (v4, corregido)
    # ===================================================================
    print("\n[3/7] Metodologia por Objetivos...")
    tipo_idx = find_first(doc, "Tipo y enfoque de investigacion")
    disenio_idx = find_first(doc, "Diseno de investigacion")
    
    if tipo_idx is not None and disenio_idx is not None:
        # Insertar ANTES de "Diseno de investigacion"
        # addprevious inserta INMEDIATAMENTE antes del target, preservando orden
        lines = [l.strip() for l in METODOLOGIA_OBJETIVOS.strip().split('\n') if l.strip()]
        disenio_element = doc.paragraphs[disenio_idx]._element
        for line in lines:
            p_element = create_new_paragraph_element(line, 'Normal')
            disenio_element.addprevious(p_element)
        total_changes += len(lines)
        print(f"  Insertados {len(lines)} parrafos de metodologia por objetivos")
    else:
        print(f"  ERROR: Tipo={tipo_idx}, Disenio={disenio_idx}")

    # ===================================================================
    # 4. IMPACTOS (v2)
    # ===================================================================
    print("\n[4/7] Impactos...")
    impact_mapping = {
        "Impactos en Ciencia y Tecnologia": "ciencia",
        "Impactos economicos": "economicos",
        "Impactos sociales": "sociales",
        "Impactos ambientales": "ambientales"
    }
    for title, key in impact_mapping.items():
        idx = find_first(doc, title)
        if idx is not None:
            for j in range(idx + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    p = doc.paragraphs[j]
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = IMPACTOS[key]
                    else:
                        p.add_run(IMPACTOS[key])
                    total_changes += 1
                    print(f"  '{title}' actualizado")
                    break
        else:
            print(f"  '{title}' no encontrado")

    # ===================================================================
    # 5. LOCALIZACION (v2)
    # ===================================================================
    print("\n[5/7] Localizacion...")
    loc_idx = find_first(doc, "Localizacion del proyecto")
    if loc_idx is not None:
        for j in range(loc_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.strip():
                p = doc.paragraphs[j]
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = LOCALIZACION
                else:
                    p.add_run(LOCALIZACION)
                total_changes += 1
                print("  Localizacion actualizada")
                break

    # ===================================================================
    # 6. RECURSOS (v4)
    # ===================================================================
    print("\n[6/7] Recursos...")
    rec_idx = find_first(doc, "Recursos necesarios")
    if rec_idx is not None:
        # Encontrar fin de seccion
        rec_end = len(doc.paragraphs)
        for j in range(rec_idx + 1, len(doc.paragraphs)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ''
            if doc.paragraphs[j].text.strip() and ('Cuerpo' in style or 'Heading' in style):
                if normalize('Localizacion') not in normalize(doc.paragraphs[j].text or ''):
                    rec_end = j
                    break
        
        # Eliminar contenido viejo
        clear_range(doc, rec_idx + 1, rec_end)
        # Insertar nuevo
        insert_after(doc.paragraphs[rec_idx]._element, RECURSOS)
        total_changes += len(RECURSOS)
        print(f"  Insertados {len(RECURSOS)} parrafos de recursos")

    # ===================================================================
    # 7. REFERENCIAS (v4)
    # ===================================================================
    print("\n[7/7] Referencias...")
    ref_idx = find_first(doc, "Referencias")
    if ref_idx is not None:
        # Encontrar fin de referencias viejas
        ref_end = len(doc.paragraphs)
        for j in range(ref_idx + 1, len(doc.paragraphs)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ''
            if doc.paragraphs[j].text.strip() and ('Cuerpo' in style or 'Heading' in style):
                ref_end = j
                break
        
        # Eliminar referencias viejas
        if ref_end > ref_idx + 1:
            clear_range(doc, ref_idx + 1, ref_end)
        # Insertar nuevas
        insert_after(doc.paragraphs[ref_idx]._element, REFERENCIAS)
        total_changes += len(REFERENCIAS)
        print(f"  Insertadas {len(REFERENCIAS)} referencias")

    # ===================================================================
    # 8. BARRIDO FINAL
    # ===================================================================
    print("\n[Barrido final] Buscando contenido residual...")
    residual_count = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text or ''
        for kw in RESIDUAL_KEYWORDS:
            if kw.lower() in text.lower() and text.strip():
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = ''
                else:
                    p.add_run('')
                residual_count += 1
                print(f"  Limpiado [{i}]: '{kw}'")
                break
    
    if residual_count == 0:
        print("  OK: Sin contenido residual")

    # ===================================================================
    # Guardar
    # ===================================================================
    doc.save(dest)
    print(f"\nDocumento guardado: {dest}")
    print(f"Total cambios: {total_changes}")
    
    # ===================================================================
    # VERIFICACION
    # ===================================================================
    print("\n=== VERIFICACION ===")
    
    # No residual
    ok = True
    for i, p in enumerate(doc.paragraphs):
        for kw in ["Raspberry", "IoT", "BoT-IoT", "TensorFlow Lite"]:
            if kw.lower() in (p.text or '').lower() and (p.text or '').strip():
                print(f"  ERROR Residual: [{i}] {kw}")
                ok = False
    if ok:
        print("  Sin contenido residual: OK")
    
    # 15 antecedentes
    ant_count = sum(1 for p in doc.paragraphs if any(
        a in (p.text or '') for a in ['Abdul-Hafiz', 'Rakesh Kumar', 'Mohit Kumar', 'Aribe Jr.',
        'Lipianina-Honcharenko', 'Raikwar', 'Singh & Kumar', 'Wang et al.',
        'Selvaraju', 'Rossler', 'Li et al.', 'Tolosana', 'Guarnera',
        'De Lima', 'Tariq']
    ))
    print(f"  Antecedentes encontrados: {ant_count}")
    
    # Metodologia por objetivos
    met_count = sum(1 for p in doc.paragraphs if 'Objetivo Especifico' in (p.text or ''))
    print(f"  Parrafos de Metodologia por Objetivos: {met_count}")
    print(f"  (Deben ser 10: 5 titulos + 5 descripciones)")
    
    # Referencias
    ref_count = sum(1 for p in doc.paragraphs if any(
        a in (p.text or '') for a in ['Arrieta, A. B.', 'Brooke, J.', 'Flach, P.', 
        'Goodfellow, I.', 'Hernandez-Sampieri', 'Naupas, H.',
        'Pan, S. J.', 'Selvaraju, R. R.', 'Sokolova, M.', 'Tukey, J. W.']
    ))
    print(f"  Referencias encontradas: {ref_count} (deben ser 10)")
    
    return dest


if __name__ == "__main__":
    path = create_document()
    print(f"\nArchivo: {path}")
