"""
Script de limpieza: reemplaza TODO el contenido residual de phishing/TabTransformer
del documento PARTE2_TRABAJO_PROYECTO_COMPLETADO.docx
"""

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import copy
import sys

try:
    sys.stdout.reconfigure(encoding='utf-8')
except:
    pass

DOC_PATH = "D:/IX/detectorIA/documentacion/PARTE2_TRABAJO_PROYECTO_COMPLETADO.docx"

# ============================================================
# MAPA DE REEMPLAZOS: viejo texto exacto -> nuevo texto
# ============================================================
REPLACEMENTS = [
    # ---------- TITULO ----------
    (
        "CLASIFICACIÓN DE URLS MALICIOSAS UTILIZANDO TABTRANSFORMER PARA LA DETECCIÓN DE ATAQUES DE PHISHING",
        "DETECCIÓN DE DEEPFAKES EN IMÁGENES MEDIANTE REDES NEURONALES CONVOLUCIONALES CON TRANSFER LEARNING Y VISUALIZACIÓN CON GRAD-CAM"
    ),
    # ---------- OBJETIVO GENERAL ----------
    (
        "Determinar la eficacia del modelo TabTransformer en la clasificación de URL maliciosas para la detección de ataques de phishing, mediante su implementación y evaluación sobre el conjunto de datos Malicious URL Detection Dataset Enhanced 2026.",
        "Desarrollar un sistema de detección de deepfakes en imágenes mediante redes neuronales convolucionales con Transfer Learning e Inteligencia Artificial Explicable basada en Grad-CAM, que permita identificar imágenes manipuladas digitalmente y generar explicaciones visuales e interpretables sobre las decisiones del modelo."
    ),
    # ---------- OBJETIVOS ESPECIFICOS (cada uno) ----------
    (
        "▪ Examinar las variables disponibles en el conjunto de datos Malicious URL Detection Dataset Enhanced 2026, identificando las características léxicas y estructurales más relevantes para diferenciar entre URL maliciosas y legítimas asociadas a ataques",
        "▪ Analizar las características presentes en imágenes reales y deepfakes mediante técnicas de procesamiento digital de imágenes para identificar patrones relevantes para su clasificación."
    ),
    (
        "▪ Elaborar la configuración del modelo TabTransformer, definiendo el tratamiento de las variables categóricas y numéricas, la cantidad de capas de atención y los demás parámetros necesarios para su aplicación a la clasificación de URL maliciosas.",
        "▪ Implementar una red neuronal convolucional basada en Transfer Learning (DenseNet-121) para la clasificación automática de imágenes reales y manipuladas digitalmente."
    ),
    (
        "▪ Desarrollar el proceso de entrenamiento del modelo TabTransformer utilizando el conjunto de datos preprocesado, aplicando las técnicas de codificación, normalización y balanceo de clases descritas en el marco teórico.",
        "▪ Evaluar el desempeño del modelo mediante métricas de precisión, recall, F1-Score y AUC para determinar su efectividad en la detección de deepfakes."
    ),
    (
        "▪ Medir el desempeño del modelo TabTransformer entrenado a partir de las métricas de exactitud, precisión, sensibilidad, puntaje F1 y área bajo la curva ROC, con el fin de determinar su capacidad para detectar URL asociadas a ataques de phishing.",
        "▪ Aplicar técnicas de Inteligencia Artificial Explicable mediante Grad-CAM para generar explicaciones visuales e interpretables sobre las decisiones tomadas por el modelo."
    ),
    (
        "▪ Contrastar los resultados alcanzados por el modelo TabTransformer con los valores reportados en los antecedentes de investigación revisados, con el propósito de situar el aporte del presente trabajo dentro del estado actual del conocimiento sobre d",
        "▪ Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificación de contenido digital manipulado."
    ),
    # ---------- FASE 1 - Objetivo y Proposito ----------
    (
        "Examinar las variables disponibles en el conjunto de datos Malicious URL Detection Dataset Enhanced 2026, identificando las características léxicas y estructurales más relevantes para diferenciar entre URLs maliciosas y legítimas asociadas a ataques ",
        "Analizar las características presentes en imágenes reales y deepfakes mediante técnicas de procesamiento digital de imágenes para identificar patrones relevantes para su clasificación."
    ),
    (
        "Conocer la estructura del conjunto de datos, evaluar su calidad e identificar las variables que aportan mayor información para reconocer URLs de phishing.",
        "Conocer la estructura del dataset, evaluar su calidad e identificar las características visuales que permiten diferenciar imágenes reales de imágenes generadas mediante inteligencia artificial."
    ),
    # ---------- FASE 2 - Objetivo y Proposito ----------
    (
        "Elaborar la configuración del modelo TabTransformer, definiendo el tratamiento de las variables categóricas y numéricas, la cantidad de capas de atención y los demás parámetros necesarios para su aplicación a la clasificación de URLs maliciosas.",
        "Implementar una red neuronal convolucional basada en Transfer Learning (DenseNet-121) para la clasificación automática de imágenes reales y manipuladas digitalmente."
    ),
    (
        "Diseñar la arquitectura del modelo que procesará simultáneamente las variables categóricas y numéricas del conjunto de datos.",
        "Diseñar la arquitectura del modelo de deep learning que aprenderá a diferenciar imágenes reales de deepfakes, aprovechando el conocimiento preentrenado en ImageNet."
    ),
    # ---------- Texto residual "El archivo explica que TabTransformer..." ----------
    (
        "El archivo explica que TabTransformer convierte las variables categóricas en embeddings, las procesa mediante bloques Transformer y luego las combina con variables numéricas para realizar la clasificación final.",
        "El modelo DenseNet-121 utiliza conexiones densas entre capas, lo que permite un mejor flujo del gradiente y una mayor eficiencia en el uso de parámetros. Cada capa recibe como entrada las características de todas las capas anteriores."
    ),
    # ---------- Algoritmos residuales ----------
    (
        "TabTransformer.\nAutoatención.\nMulti-Head Attention.",
        "DenseNet-121.\nAutoatención.\nTransfer Learning."
    ),
    # ---------- FASE 3 - Objetivo y Proposito ----------
    (
        "Desarrollar el proceso de entrenamiento del modelo TabTransformer utilizando el conjunto de datos preprocesado, aplicando técnicas de codificación, normalización y balanceo de clases.",
        "Implementar una red neuronal convolucional basada en Transfer Learning (DenseNet-121) para la clasificación automática de imágenes reales y manipuladas digitalmente."
    ),
    (
        "Preparar los datos y entrenar el modelo para que aprenda a diferenciar URLs legítimas de URLs relacionadas con ataques de phishing.",
        "Preparar los datos y entrenar el modelo DenseNet-121 para que aprenda a diferenciar imágenes reales de deepfakes, utilizando transfer learning y técnicas de regularización."
    ),
    # ---------- FASE 4 - Objetivo y Proposito ----------
    (
        "Medir el desempeño del modelo TabTransformer entrenado a partir de las métricas de exactitud, precisión, sensibilidad, puntaje F1 y área bajo la curva ROC, con el fin de determinar su capacidad para detectar URLs asociadas a ataques de phishing.",
        "Evaluar el desempeño del modelo mediante métricas de precisión, recall, F1-Score y AUC para determinar su efectividad en la detección de deepfakes."
    ),
    (
        "Determinar objetivamente si el modelo identifica correctamente las URLs de phishing y si mantiene un nivel aceptable de errores.",
        "Determinar objetivamente si el modelo identifica correctamente las imágenes deepfake y mantiene un nivel aceptable de errores en condiciones normales y adversarias."
    ),
    # ---------- FASE 5 - Objetivo y Relacion con objetivo general ----------
    (
        "Contrastar los resultados alcanzados por el modelo TabTransformer con los valores reportados en los antecedentes de investigación revisados, con el propósito de situar el aporte del presente trabajo dentro del estado actual del conocimiento sobre det",
        "Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificación de contenido digital manipulado."
    ),
    (
        "Esta fase permite cumplir directamente el objetivo general, porque con los resultados obtenidos se determinará la eficacia real del modelo TabTransformer.",
        "Esta fase permite cumplir directamente el objetivo general, porque con los resultados obtenidos se determina la eficacia real del modelo DenseNet-121 en la detección explicable de deepfakes."
    ),
    # ---------- TabTransformer en Algoritmos ----------
    (
        "TabTransformer.\nAnálisis comparativo.\nEvaluación experimental.",
        "DenseNet-121.\nTransfer Learning.\nGrad-CAM.\nAnálisis comparativo.\nEvaluación experimental."
    ),
]


def set_paragraph_text(paragraph, text):
    """Replace text in a paragraph preserving its first run."""
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def clean_document():
    doc = Document(DOC_PATH)
    
    # Para tracking
    replaced_count = 0
    not_found = []
    
    for old_text, new_text in REPLACEMENTS:
        found = False
        for i, p in enumerate(doc.paragraphs):
            if old_text in (p.text or ''):
                set_paragraph_text(p, new_text)
                replaced_count += 1
                found = True
                print(f"  OK [{i}]: Reemplazado texto que contenia: '{old_text[:60]}...'")
                break
        if not found:
            not_found.append(old_text[:80])
    
    # Guardar
    doc.save(DOC_PATH)
    
    print(f"\nResumen:")
    print(f"  Reemplazos exitosos: {replaced_count}")
    print(f"  No encontrados: {len(not_found)}")
    
    if not_found:
        print("\nNO encontrados:")
        for nf in not_found:
            print(f"  - {nf}...")
    
    # Verificacion final
    print("\n=== VERIFICACION FINAL ===")
    keywords_final = ['phishing', 'TabTransformer', 'URL malicios', 'URLs legítimas']
    still_bad = []
    for p in doc.paragraphs:
        txt = p.text or ''
        for kw in keywords_final:
            if kw.lower() in txt.lower():
                still_bad.append((p.style.name, txt[:120]))
                break
    
    if still_bad:
        print(f"  AUN quedan {len(still_bad)} parrafos con contenido residual:")
        for style, txt in still_bad:
            print(f"    [{style}] {txt}")
    else:
        print("  NO HAY contenido residual de phishing! TODO LIMPIO!")
    
    return len(still_bad) == 0


if __name__ == "__main__":
    ok = clean_document()
    exit(0 if ok else 1)
