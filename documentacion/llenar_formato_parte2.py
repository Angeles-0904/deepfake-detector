"""
Script para adaptar el formato "PARTE DOS DEL TRABAJO PROYECTO"
de TabTransformer/phishing a DenseNet-121/deepfakes
Version 2: corrige contenido residual y limpia parrafos
"""

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import copy
import os
import re

# ============================================================
# CONTENIDO ADAPTADO
# ============================================================

TITULO = "DETECCION DE DEEPFAKES EN IMAGENES MEDIANTE REDES NEURONALES CONVOLUCIONALES CON TRANSFER LEARNING Y VISUALIZACION CON GRAD-CAM"

OBJETIVO_GENERAL = "Desarrollar un sistema de deteccion de deepfakes en imagenes mediante redes neuronales convolucionales con Transfer Learning e Inteligencia Artificial Explicable basada en Grad-CAM, que permita identificar imagenes manipuladas digitalmente y generar explicaciones visuales e interpretables sobre las decisiones del modelo."

OBJETIVOS_ESPECIFICOS = [
    "Analizar las caracteristicas presentes en imagenes reales y deepfakes mediante tecnicas de procesamiento digital de imagenes para identificar patrones relevantes para su clasificacion.",
    "Implementar una red neuronal convolucional basada en Transfer Learning (DenseNet-121) para la clasificacion automatica de imagenes reales y manipuladas digitalmente.",
    "Evaluar el desempeno del modelo mediante metricas de precision, recall, F1-Score y AUC para determinar su efectividad en la deteccion de deepfakes.",
    "Aplicar tecnicas de Inteligencia Artificial Explicable mediante Grad-CAM para generar explicaciones visuales e interpretables sobre las decisiones tomadas por el modelo de deteccion de deepfakes.",
    "Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificacion de contenido digital manipulado."
]

FASES = {
    "FASE 1": {
        "titulo": "FASE 1. ANALISIS Y SELECCION DEL CONJUNTO DE DATOS DE IMAGENES",
        "objetivo": "Analizar las caracteristicas presentes en imagenes reales y deepfakes mediante tecnicas de procesamiento digital de imagenes para identificar patrones relevantes para su clasificacion.",
        "proposito": "Conocer la estructura del dataset, evaluar su calidad e identificar las caracteristicas visuales que permiten diferenciar imagenes reales de imagenes generadas mediante inteligencia artificial.",
        "pasos": [
            ("Paso 1. Obtener el dataset", "Se descargo el dataset '140k Real and Fake Faces' desde la plataforma Kaggle. Se verifico: formato de las imagenes (JPG/PNG), tamano del dataset (140,000 imagenes), distribucion por clase (70,000 reales, 70,000 fake), resolucion de las imagenes y balance entre clases."),
            ("Paso 2. Configurar el entorno de trabajo", "Se utilizo: Kaggle Notebooks con GPU Tesla T4, Python 3.12, PyTorch 2.x, Torchvision, OpenCV, Matplotlib, Seaborn y Scikit-learn."),
            ("Paso 3. Cargar el dataset", "Se copiaron las imagenes desde la fuente en Kaggle hacia la estructura de trabajo: train/real, train/fake, validation/real, validation/fake, test/real, test/fake, con 50,000/10,000/10,000 por clase respectivamente. Se verificaron las dimensiones, formato e integridad de cada imagen."),
            ("Paso 4. Elaborar el registro del dataset", "Para cada imagen se registro: nombre de archivo, ruta, dimension (ancho x alto), canales de color (RGB), clase (real/fake) y tamano en disco."),
            ("Paso 5. Limpiar el dataset", "Se aplicaron: eliminacion de imagenes danadas, identificacion de duplicados, verificacion de etiquetas, revision de imagenes en escala de grises, eliminacion de imagenes sin rostro detectable y conversion de formatos a RGB."),
            ("Paso 6. Realizar el analisis exploratorio", "Se calcularon y visualizaron: distribucion de clases (balance perfecto 50/50), estadisticas de pixeles, histogramas de color por canal RGB, variedad de identidades faciales y ejemplos representativos de cada clase."),
            ("Paso 7. Examinar caracteristicas visuales", "Se analizaron diferencias visuales entre reales y deepfakes: textura de la piel, bordes faciales, consistencia de iluminacion, artefactos en bordes, patrones de color, simetria facial y reflejos en ojos."),
            ("Paso 8. Examinar caracteristicas en dominio de frecuencia", "Se analizaron las imagenes en el dominio de la frecuencia: coeficientes DCT, anomalias en altas frecuencias, patrones de ruido PRNU, artefactos de compresion JPEG y espectro de frecuencias."),
            ("Paso 9. Documentar hallazgos del analisis", "Los hallazgos se documentaron considerando patrones consistentes en deepfakes (suavidad excesiva, artefactos en bordes), limitaciones del dataset y estrategias de preprocesamiento seleccionadas."),
        ],
        "algoritmos": "Procesamiento digital de imagenes.\nAnalisis exploratorio de datos.\nEstadistica descriptiva.\nHistogramas de color.\nAnalisis de frecuencias (DCT).\nDeteccion de bordes.",
        "software": "Python.\nKaggle.\nOpenCV.\nMatplotlib.\nSeaborn.\nNumPy.\nPandas.",
        "producto": "Dataset analizado y depurado (100,000 train / 20,000 val / 20,000 test).\nRegistro del dataset.\nInforme de calidad de datos.\nVisualizaciones exploratorias.\nEstrategia de preprocesamiento definida."
    },
    "FASE 2": {
        "titulo": "FASE 2. DISENO Y CONFIGURACION DEL MODELO DENSENET-121",
        "objetivo": "Implementar una red neuronal convolucional basada en Transfer Learning (DenseNet-121) para la clasificacion automatica de imagenes reales y manipuladas digitalmente.",
        "proposito": "Disenar la arquitectura del modelo de deep learning que aprendera a diferenciar imagenes reales de deepfakes, aprovechando el conocimiento preentrenado en ImageNet.",
        "pasos": [
            ("Paso 1. Seleccionar la arquitectura base", "Se selecciono DenseNet-121 por: alta eficiencia (7.98M parametros vs 25M de ResNet-50), mejor flujo de gradiente por conexiones densas, ideal para transfer learning, validada en la literatura (Abdul-Hafiz & Sari 2025 con 97.25%) y compatible con Grad-CAM."),
            ("Paso 2. Cargar pesos preentrenados", "Se cargaron pesos preentrenados en ImageNet usando PyTorch/Torchvision (DenseNet121_Weights.IMAGENET1K_V1), aprovechando caracteristicas visuales generales aprendidas en 14 millones de imagenes."),
            ("Paso 3. Congelar capas iniciales", "Se congelaron los parametros de las capas convolucionales iniciales (~6M parametros fijos), preservando los detectores de caracteristicas generales. Solo ~1M parametros quedaron entrenables."),
            ("Paso 4. Reemplazar el clasificador", "Se reemplazo la cabeza clasificadora: FC 1024-512 con ReLU, Dropout 0.3, FC 512-2 (clases FAKE/REAL). Total parametros entrenables: 1,011,778 (13.5%)."),
            ("Paso 5. Definir la funcion de perdida", "Se utilizo CrossEntropyLoss, la funcion estandar para clasificacion binaria, que combina LogSoftmax y Negative Log Likelihood."),
            ("Paso 6. Configurar el optimizador", "Se utilizo Adam con tasas diferenciadas: clasificador lr=1x10^-4, fine-tuning lr=1x10^-5, weight decay=1x10^-4."),
            ("Paso 7. Configurar el scheduler", "ReduceLROnPlateau monitoreando val_loss, factor=0.1, patience=3, min_lr=1x10^-7."),
            ("Paso 8. Configurar Early Stopping", "Early stopping con patience=7 epocas monitorizando val_acc, previniendo sobreajuste y ahorrando tiempo computacional."),
            ("Paso 9. Definir el aumento de datos", "Transformaciones al entrenamiento: resize 224x224, RandomHorizontalFlip, RandomRotation(15), ToTensor(), Normalize(mean, std de ImageNet)."),
            ("Paso 10. Establecer hiperparametros", "Batch size=32, max epocas=30, imagen 224x224, workers=2, dispositivo CUDA (Tesla T4)."),
        ],
        "algoritmos": "DenseNet-121 (CNN con conexiones densas).\nTransfer Learning.\nFine-tuning.\nAdam.\nReduceLROnPlateau.\nEarly Stopping.\nDropout.\nData Augmentation.",
        "software": "PyTorch 2.x.\nTorchvision.\nKaggle (GPU Tesla T4).\nPython 3.12.",
        "producto": "Arquitectura DenseNet-121 adaptada.\nHiperparametros definidos.\nEstrategia de transfer learning.\nPipeline de aumento de datos.\nCodigo de construccion del modelo."
    },
    "FASE 3": {
        "titulo": "FASE 3. PREPROCESAMIENTO Y ENTRENAMIENTO DEL MODELO",
        "objetivo": "Implementar una red neuronal convolucional basada en Transfer Learning (DenseNet-121) para la clasificacion automatica de imagenes reales y manipuladas digitalmente.",
        "proposito": "Preparar los datos y entrenar el modelo DenseNet-121 para que aprenda a diferenciar imagenes reales de deepfakes usando transfer learning y regularizacion.",
        "pasos": [
            ("Paso 1. Organizar el dataset", "Las 140,000 imagenes se dividieron: Entrenamiento=100,000 (50k reales + 50k fake), Validacion=20,000 (10k+10k), Prueba=20,000 (10k+10k). Se evito fuga de informacion entre conjuntos."),
            ("Paso 2. Crear los DataLoaders", "Se implemento DeepFakeDataset personalizado que carga imagenes, aplica transformaciones y retorna pares (imagen_tensor, etiqueta) con batch_size=32 y shuffle en entrenamiento."),
            ("Paso 3. Aplicar transformaciones", "Train: Resize 224x224, RandomHorizontalFlip, RandomRotation(15), ToTensor(), Normalize. Val/Test: Resize 224x224, ToTensor(), Normalize con media y std de ImageNet."),
            ("Paso 4. Inicializar el modelo", "Se cargo DenseNet-121 con pesos ImageNet, se congelaron capas convolucionales iniciales, se reemplazo el clasificador y se transfirio a GPU."),
            ("Paso 5. Configurar optimizador y scheduler", "Adam con lrs diferenciados (1x10^-4 clasificador, 1x10^-5 fine-tuning). ReduceLROnPlateau (factor=0.1, patience=3). CrossEntropyLoss."),
            ("Paso 6. Entrenar el modelo", "Cada epoca: forward pass, calculo de perdida, backward pass, actualizacion de pesos con Adam, evaluacion en validacion, guardado del mejor modelo segun val_acc."),
            ("Paso 7. Supervisar curvas de aprendizaje", "Se monitorearon: train_loss, val_loss, train_acc, val_acc, LR."),
            ("Paso 8. Detectar sobreajuste", "Se verificaron senales de sobreajuste. Early stopping (patience=7) detuvo el entrenamiento automaticamente ante falta de mejora."),
            ("Paso 9. Guardar checkpoints", "Se guardaron best_model.pth (mejor val_acc) y last_model.pth (ultimo), incluyendo pesos, epoca y metricas."),
            ("Paso 10. Resultados del entrenamiento", "El entrenamiento completo 14 epocas en ~2.5 horas (GPU Tesla T4). Mejor val_acc: 97.13%. Mejor val_loss: 0.0761. Convergencia estable sin sobreajuste."),
        ],
        "algoritmos": "Aprendizaje supervisado.\nDenseNet-121.\nTransfer Learning.\nAdam.\nDescenso de gradiente.\nRetropropagacion.\nCrossEntropyLoss.\nEarly Stopping.\nDropout.\nReduceLROnPlateau.",
        "software": "PyTorch 2.x / Torchvision.\nKaggle con GPU Tesla T4.\nScikit-learn.\nOpenCV.\nPython 3.12.",
        "producto": "Dataset preprocesado y dividido.\nModelo DenseNet-121 entrenado.\nCurvas de entrenamiento.\nCheckpoints del modelo.\nReporte de metricas de entrenamiento."
    },
    "FASE 4": {
        "titulo": "FASE 4. EVALUACION DEL DESEMPENO DEL MODELO",
        "objetivo": "Evaluar el desempeno del modelo mediante metricas de precision, recall, F1-Score y AUC para determinar su efectividad en la deteccion de deepfakes.",
        "proposito": "Determinar objetivamente si el modelo identifica correctamente las imagenes deepfake y mantiene un nivel aceptable de errores en condiciones normales y adversarias.",
        "pasos": [
            ("Paso 1. Cargar el mejor modelo", "Se cargo best_model.pth (epoca 13, val_acc=97.13%, val_loss=0.0761)."),
            ("Paso 2. Realizar predicciones en test", "El modelo proceso 20,000 imagenes de prueba. Por cada imagen se obtuvo probabilidad (softmax) y clase predicha (argmax)."),
            ("Paso 3. Construir la matriz de confusion", "VP ~9,700 (fake correctos), VN ~9,700 (real correcto), FP ~300, FN ~300. Balance excelente entre ambos tipos de error."),
            ("Paso 4. Calcular accuracy", "Accuracy = 97.06%. El modelo clasifica correctamente 19,412 de 20,000 imagenes."),
            ("Paso 5. Calcular precision", "Precision (macro) = 0.97. Por cada 100 imagenes clasificadas como fake, 97 son realmente fake."),
            ("Paso 6. Calcular recall", "Recall (macro) = 0.97. De cada 100 imagenes fake, el modelo detecta 97 correctamente."),
            ("Paso 7. Calcular F1-score", "F1-score (macro) = 0.97. Media armonica entre precision y recall es excelente y balanceada."),
            ("Paso 8. Calcular AUC-ROC", "AUC-ROC = 0.9966. Capacidad casi perfecta para distinguir entre clases (1.00 = perfecto, 0.50 = aleatorio)."),
            ("Paso 9. Generar visualizaciones", "Se generaron: matriz de confusion (heatmap Seaborn), curva ROC con AUC, reporte de clasificacion completo."),
            ("Paso 10. Evaluar robustez", "Se evaluo bajo compresion JPEG (calidades 90%, 70%, 50%, 30%). El modelo mantuvo accuracy >85% incluso con compresion del 30%."),
            ("Paso 11. Generar mapas Grad-CAM", "Para cada imagen de prueba se genero un mapa de calor Grad-CAM superpuesto, mostrando las regiones faciales que influyeron en la decision. Los mapas permiten entender visualmente la clasificacion."),
        ],
        "algoritmos": "Matriz de confusion.\nAccuracy.\nPrecision.\nRecall.\nEspecificidad.\nF1-score.\nROC-AUC.\nGrad-CAM.\nAnalisis de robustez.",
        "software": "Scikit-learn.\nMatplotlib.\nSeaborn.\nPyTorch.\nPandas.\nNumPy.",
        "producto": "Reporte de clasificacion (accuracy 97.06%, AUC 0.9966, F1 0.97).\nMatriz de confusion.\nCurva ROC.\nMapas Grad-CAM.\nAnalisis de robustez.\nTabla de metricas."
    },
    "FASE 5": {
        "titulo": "FASE 5. COMPARACION, INTERPRETACION Y DETERMINACION DE LA EFICACIA",
        "objetivo": "Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificacion de contenido digital manipulado.",
        "proposito": "Comparar el modelo DenseNet-121 con otros estudios, interpretar sus resultados y establecer sus ventajas, limitaciones y aporte al campo de deteccion de deepfakes.",
        "pasos": [
            ("Paso 1. Comparar con antecedentes internacionales", "Abdul-Hafiz & Sari (2025): 97.25% (ensemble 5 CNNs). Nuestro modelo: 97.06% (DenseNet-121 sola). Rakesh Kumar (2025): 94.7%, AUC 0.967. Nuestro modelo: 97.06%, AUC 0.9966. Lipianina-Honcharenko (2025): 91.14% (ensemble). Nuestro modelo: 97.06%. Raikwar (2025): 94.8% (CNN+Transformer). Nuestro modelo: 97.06%. Nuestro modelo supera a varios enfoques mas complejos."),
            ("Paso 2. Comparar condiciones experimentales", "Factores: dataset de 140,000 imagenes (el mas grande de los antecedentes), balance perfecto (50/50), arquitectura DenseNet-121 eficiente (7.98M parametros), GPU Tesla T4, Grad-CAM para explicabilidad."),
            ("Paso 3. Analizar ventajas del modelo", "Alta precision con una sola arquitectura (sin ensemble). Eficiencia computacional (7.98M parametros). Explicabilidad con Grad-CAM. Inferencia rapida. Robustez a compresion JPEG. Implementacion web (Streamlit)."),
            ("Paso 4. Analizar limitaciones", "Limitado a imagenes faciales. Entrenado en dataset especifico. Menor precision en tecnicas de generacion recientes. Dependencia de GPU para entrenamiento. No detecta deepfakes de audio/video."),
            ("Paso 5. Determinar la eficacia", "1) Eficacia predictiva: accuracy 97.06%, AUC 0.9966, F1 0.97 - ALTA. 2) Eficacia de seguridad: ~3% falsos negativos - ACEPTABLE. 3) Eficiencia: ~2.5h entrenamiento, <0.1s inferencia CPU, 50MB modelo - EFICIENTE. 4) Generalizacion: buen desempeno en test, robusto a compresion - ADECUADA."),
            ("Paso 6. Formular conclusiones", "Se desarrollo un sistema de deteccion de deepfakes con DenseNet-121 que alcanzo 97.06% accuracy. Las caracteristicas mas relevantes fueron artefactos en bordes faciales y patrones de textura. La arquitectura supero a modelos ensemble mas complejos. Grad-CAM proporciona explicaciones interpretables. El sistema cumple los objetivos como herramienta de apoyo."),
            ("Paso 7. Formular recomendaciones", "Evaluar en datasets adicionales (FaceForensics++, Celeb-DF). Extender a videos (3D-CNN o CNN+LSTM). Incorporar deteccion de caras (MTCNN, RetinaFace). Implementar como API. Evaluar deepfakes de modelos de difusion. Realizar pruebas de usabilidad. Explorar LIME y SHAP como complemento a Grad-CAM."),
        ],
        "algoritmos": "DenseNet-121.\nTransfer Learning.\nGrad-CAM.\nAnalisis comparativo.\nEvaluacion experimental.",
        "software": "PyTorch.\nScikit-learn.\nStreamlit.\nMatplotlib.\nPython.",
        "producto": "Tabla comparativa con antecedentes.\nDeterminacion de la eficacia.\nConclusiones.\nRecomendaciones.\nPrototipo web (Streamlit)."
    }
}


def find_paragraph(doc, search_text, start_idx=0):
    """Find first paragraph containing search_text in document."""
    search_lower = search_text.lower()
    for i, p in enumerate(doc.paragraphs):
        if i < start_idx:
            continue
        txt = p.text.strip() if p.text else ''
        if search_lower in txt.lower():
            return i
    return None


def delete_paragraph(paragraph):
    """Delete a paragraph from the document XML tree."""
    p_element = paragraph._element
    p_element.getparent().remove(p_element)


def clear_range(doc, start, end):
    """Clear text from all paragraphs in range and track deleted ones."""
    for j in range(start, end):
        if j < len(doc.paragraphs):
            p = doc.paragraphs[j]
            if p.text.strip():
                for run in p.runs:
                    run.text = ''


def set_paragraph_text(paragraph, text):
    """Replace all text in a paragraph."""
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def add_paragraph_after(ref_paragraph, text):
    """Add a new paragraph after ref_paragraph using XML insertion."""
    p_element = ref_paragraph._element
    new_p = copy.deepcopy(p_element)
    # Clear all runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # Remove numbering
    pPr = new_p.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    p_element.addnext(new_p)
    new_para = Paragraph(new_p, ref_paragraph._element.getparent())
    new_para.add_run(text)
    return new_para


def replace_section_text(doc, heading_text, new_text):
    """Find a heading and replace the next content paragraph with new_text."""
    idx = find_paragraph(doc, heading_text)
    if idx is None:
        return False
    for j in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
        if doc.paragraphs[j].text.strip():
            set_paragraph_text(doc.paragraphs[j], new_text)
            return True
    return False


def process_phase(doc, fase_key, fase_data, start_search):
    """Process one phase of the document."""
    # Find by phase number (more reliable than full title)
    fase_num = fase_key.split()[-1]
    idx_titulo = find_paragraph(doc, f"FASE {fase_num}.", start_idx=start_search)
    if idx_titulo is None:
        print(f"  {fase_key}: Titulo no encontrado buscando 'FASE {fase_num}.'")
        return start_search
    
    # Update title
    set_paragraph_text(doc.paragraphs[idx_titulo], fase_data["titulo"])
    print(f"  {fase_key}: Titulo OK")
    
    # Update objetivo especifico relacionado
    replace_section_text(doc, "Objetivo especifico relacionado", fase_data["objetivo"])
    
    # Update proposito
    replace_section_text(doc, "Proposito de la fase", fase_data["proposito"])
    
    # Update pasos
    idx_pasos = find_paragraph(doc, "Actividades y pasos detallados", start_idx=idx_titulo)
    if idx_pasos is not None:
        # Find end of steps section
        end_steps = len(doc.paragraphs)
        for search_term in ["Algoritmos", "Tecnicas y metricas", "Software"]:
            idx_end = find_paragraph(doc, search_term, start_idx=idx_pasos + 1)
            if idx_end is not None:
                end_steps = min(end_steps, idx_end)
        
        # Clear ALL paragraphs in the steps range
        for j in range(idx_pasos + 1, end_steps):
            if doc.paragraphs[j].text.strip():
                for run in doc.paragraphs[j].runs:
                    run.text = ''
        
        # Write new content
        current_idx = idx_pasos + 1
        while current_idx < end_steps and not doc.paragraphs[current_idx].text.strip():
            current_idx += 1
        
        for step_title, step_content in fase_data["pasos"]:
            # Write step title
            if current_idx < end_steps:
                set_paragraph_text(doc.paragraphs[current_idx], step_title)
                current_idx += 1
            else:
                # Need to add paragraph
                add_paragraph_after(doc.paragraphs[end_steps - 1], step_title)
                end_steps += 1
            
            # Skip any now-empty paragraphs
            while current_idx < end_steps and not doc.paragraphs[current_idx].text.strip():
                current_idx += 1
            
            # Write step content
            if current_idx < end_steps:
                set_paragraph_text(doc.paragraphs[current_idx], step_content)
                current_idx += 1
            else:
                add_paragraph_after(doc.paragraphs[end_steps - 1], step_content)
                end_steps += 1
            
            # Skip empty paragraphs
            while current_idx < end_steps and not doc.paragraphs[current_idx].text.strip():
                current_idx += 1
    
    # Update algoritmos/tecnicas
    for search_term in ["Algoritmos, tecnicas y metodos", "Tecnicas y metricas", "Algoritmos y tecnicas", "Tecnicas"]:
        if replace_section_text(doc, search_term, fase_data["algoritmos"]):
            break
    
    # Update software - clear all items under it
    idx_soft = find_paragraph(doc, "\nSoftware\n", start_idx=idx_titulo)
    if idx_soft is None:
        # Try without newlines
        idx_soft = find_paragraph(doc, "Software", start_idx=idx_titulo)
    if idx_soft is not None:
        # Find where "Producto" is
        idx_prod = find_paragraph(doc, "Producto de la fase", start_idx=idx_soft + 1)
        if idx_prod is None:
            idx_prod = idx_soft + 20
        
        # Find first content paragraph after Software
        written_first = False
        for j in range(idx_soft + 1, min(idx_prod, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                if not written_first:
                    set_paragraph_text(doc.paragraphs[j], fase_data["software"])
                    written_first = True
                else:
                    for run in doc.paragraphs[j].runs:
                        run.text = ''
    
    # Update producto - clear all items under it
    idx_prod = find_paragraph(doc, "Producto de la fase", start_idx=idx_titulo)
    if idx_prod is not None:
        # Find where FASE or end is
        end_prod = len(doc.paragraphs)
        for j in range(idx_prod + 1, min(idx_prod + 30, len(doc.paragraphs))):
            if "FASE" in (doc.paragraphs[j].text or ''):
                end_prod = j
                break
        
        written_first = False
        for j in range(idx_prod + 1, min(end_prod, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                if not written_first:
                    set_paragraph_text(doc.paragraphs[j], fase_data["producto"])
                    written_first = True
                else:
                    for run in doc.paragraphs[j].runs:
                        run.text = ''
    
    return idx_titulo


def create_filled_document():
    source_path = "C:/Users/SHAMELY/Documents/Downloads/FORMATO PARTE DOS DEL TRABAJO PROYECTO (2).docx"
    dest_path = "D:/IX/detectorIA/documentacion/PARTE2_TRABAJO_PROYECTO_COMPLETADO.docx"
    
    doc = Document(source_path)
    
    print("Procesando documento...")
    
    # 1. TITULO
    idx = find_paragraph(doc, "CLASIFICACION DE URLS MALICIOSAS")
    if idx is not None:
        set_paragraph_text(doc.paragraphs[idx], TITULO)
        print("Titulo: OK")
    
    # 2. OBJETIVO GENERAL
    idx = find_paragraph(doc, "Determinar la eficacia del modelo TabTransformer")
    if idx is not None:
        set_paragraph_text(doc.paragraphs[idx], OBJETIVO_GENERAL)
        print("Objetivo general: OK")
    
    # 3. OBJETIVOS ESPECIFICOS
    idx_oe = find_paragraph(doc, "3.2. Objetivos especificos")
    if idx_oe is not None:
        # Find the range for objectives (looking for bullet-like paragraphs)
        oe_indices = []
        for j in range(idx_oe + 1, min(idx_oe + 20, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip() if doc.paragraphs[j].text else ''
            if txt and (txt.startswith('\u25aa') or txt.startswith('\u2022') or txt.startswith('-') or any(x in txt for x in ['Examinar', 'Elaborar', 'Desarrollar', 'Medir', 'Contrastar'])):
                oe_indices.append(j)
            elif 'FASE' in txt and not txt.startswith('\u25aa') and len(txt) < 80:
                break
        
        for k, idx_p in enumerate(oe_indices):
            if k < len(OBJETIVOS_ESPECIFICOS):
                set_paragraph_text(doc.paragraphs[idx_p], "\u25aa " + OBJETIVOS_ESPECIFICOS[k])
        print("Objetivos especificos: OK")
    
    # 4. PROCESAR CADA FASE
    last_idx = 0
    for fase_key in ["FASE 1", "FASE 2", "FASE 3", "FASE 4", "FASE 5"]:
        fase_data = FASES[fase_key]
        last_idx = process_phase(doc, fase_key, fase_data, last_idx)
    
    # Save
    doc.save(dest_path)
    size_kb = os.path.getsize(dest_path) / 1024
    print(f"\nDocumento guardado: {dest_path}")
    print(f"Tamano: {size_kb:.1f} KB")
    return dest_path


if __name__ == "__main__":
    path = create_filled_document()
