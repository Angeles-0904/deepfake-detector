"""
Script v4: Mejora la Metodologia con vinculacion objetivo-metodologia,
limpia contenido residual de IoT/Raspberry Pi, actualiza las Referencias.
Corrige: acentos, parrafos vacios, codigo muerto.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import unicodedata

# Normaliza texto: elimina acentos para busqueda
def normalize(text):
    if text is None:
        return ''
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii').lower().strip()


# ============================================================
# CONTENIDO MEJORADO
# ============================================================

METODOLOGIA_OBJETIVOS = """
La presente investigacion empleo una metodologia de tipo aplicada con enfoque cuantitativo y nivel explicativo-predictivo y tecnologico, desarrollada bajo un diseno experimental computacional. A continuacion se describe la relacion entre cada objetivo especifico y la metodologia empleada para su cumplimiento:

Objetivo Especifico 1: "Analizar las caracteristicas presentes en imagenes reales y deepfakes mediante tecnicas de procesamiento digital de imagenes para identificar patrones relevantes para su clasificacion."

Para el cumplimiento de este objetivo se aplico la metodologia de analisis exploratorio de datos (EDA por sus siglas en ingles). Esta metodologia, descrita por Tukey (1977) y adaptada por Hernandez-Sampieri y Mendoza (2018) como parte del enfoque cuantitativo, permitio examinar visual y estadisticamente las caracteristicas de las imagenes del dataset "140k Real and Fake Faces". Se analizaron distribuciones de intensidad de pixeles, histogramas de color en los canales RGB, niveles de ruido, frecuencia de texturas y artefactos de compresion. Se utilizaron las bibliotecas OpenCV y Matplotlib en Python para el procesamiento y visualizacion de las imagenes. Este analisis exploratorio sento las bases para disenar el preprocesamiento que se aplicaria en etapas posteriores. Segun Hernandez-Sampieri y Mendoza (2018), el analisis exploratorio es la primera etapa en investigaciones cuantitativas que buscan describir las caracteristicas de los datos antes de someterlos a pruebas mas complejas.

Objetivo Especifico 2: "Implementar una red neuronal convolucional basada en Transfer Learning para la clasificacion automatica de imagenes reales y manipuladas digitalmente."

Para este objetivo se empleo la metodologia de Transfer Learning (aprendizaje por transferencia), que consiste en reutilizar los pesos de un modelo preentrenado en una tarea base y adaptarlos a una tarea objetivo relacionada. Esta metodologia, formalizada por Pan y Yang (2010) y ampliamente adoptada en vision por computadora, permitio tomar la arquitectura DenseNet-121 preentrenada en ImageNet (con 1.2 millones de imagenes y 1000 clases) y adaptarla a la clasificacion binaria (real vs. deepfake) mediante fine-tuning. Se congelaron las capas convolucionales inferiores para preservar las caracteristicas generales (bordes, texturas, formas) y se reemplazo el clasificador por una nueva cabeza fully connected. La implementacion se realizo con PyTorch y Torchvision. Este enfoque es consistente con lo senalado por Goodfellow, Bengio y Courville (2016), quienes describen el transfer learning como una estrategia efectiva cuando los datos de la tarea objetivo son insuficientes para entrenar una red desde cero.

Objetivo Especifico 3: "Evaluar el desempeno del modelo mediante metricas de precision, recall, F1-Score y AUC para determinar su efectividad en la deteccion de deepfakes."

Se aplico la metodologia de evaluacion de clasificadores binarios mediante metricas estandar de rendimiento, conforme a los principios establecidos por Sokolova y Lapalme (2009) y Flach (2012). Esta metodologia implico: (a) calculo de la matriz de confusion (verdaderos positivos, verdaderos negativos, falsos positivos, falsos negativos); (b) derivacion de metricas derivadas: accuracy, precision, recall (sensibilidad), especificidad, F1-score; (c) generacion de la curva ROC y calculo del AUC (Area Under the Curve). La evaluacion se realizo sobre un conjunto de prueba independiente de 20,000 imagenes que el modelo no habia visto durante el entrenamiento. Segun Hernandez-Sampieri y Mendoza (2018), la evaluacion mediante metricas cuantitativas constituye un procedimiento central en investigaciones de enfoque cuantitativo y nivel explicativo-predictivo.

Objetivo Especifico 4: "Aplicar tecnicas de Inteligencia Artificial Explicable mediante Grad-CAM para generar explicaciones visuales e interpretables sobre las decisiones tomadas por el modelo de deteccion de deepfakes."

Para este objetivo se empleo la metodologia de Inteligencia Artificial Explicable (XAI, por sus siglas en ingles), especificamente la tecnica Grad-CAM (Gradient-weighted Class Activation Mapping) desarrollada por Selvaraju et al. (2017). La metodologia XAI busca hacer interpretables las decisiones de modelos de caja negra como las redes neuronales profundas. El procedimiento consistio en: (a) propagar una imagen hacia adelante a traves del modelo para obtener la prediccion; (b) calcular los gradientes de la clase predecida con respecto a los mapas de activacion de la ultima capa convolucional; (c) ponderar los mapas de activacion por los gradientes promedio; y (d) superponer el mapa de calor resultante sobre la imagen original. Segun Arrieta et al. (2020), la XAI se ha convertido en un area prioritaria de investigacion debido a la necesidad de confianza y transparencia en sistemas de inteligencia artificial.

Objetivo Especifico 5: "Validar la utilidad del sistema propuesto como herramienta de apoyo para la identificacion de contenido digital manipulado."

Se aplico la metodologia de validacion funcional y de usabilidad, basada en dos componentes. El primer componente fue la validacion funcional, que consistio en verificar que el prototipo web (implementado en Streamlit) recibiera correctamente la imagen de entrada, ejecutara el modelo de clasificacion, generara los mapas Grad-CAM y mostrara los resultados de forma clara e interactiva. El segundo componente fue la evaluacion de usabilidad, aplicando el cuestionario System Usability Scale (SUS) desarrollado por Brooke (1996), que consta de 10 preguntas con puntuacion de 1 a 5. La muestra para esta evaluacion fue no probabilistica de tipo intencional o por conveniencia, seleccionando entre 15 y 30 usuarios con conocimientos basicos en tecnologia. Segun Naupas et al. (2018), la validacion funcional y de usabilidad constituye un procedimiento propio de investigaciones de tipo aplicado y nivel tecnologico.
"""

def delete_paragraph_elements(doc, start_idx, end_idx=None):
    """Remove paragraph elements from the XML tree (not just clear text)."""
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    deleted = 0
    # Work backwards to preserve indices
    for i in range(min(end_idx, len(doc.paragraphs)) - 1, start_idx - 1, -1):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            deleted += 1
    print(f"  Eliminados {deleted} parrafos del XML (indices {start_idx} a {end_idx - 1})")


def create_new_paragraph_element(text, style='Normal'):
    """Create a <w:p> element with <w:r><w:t> containing the text."""
    p_element = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    p_element.append(pPr)
    
    r_element = OxmlElement('w:r')
    t_element = OxmlElement('w:t')
    t_element.set(qn('xml:space'), 'preserve')
    t_element.text = text
    r_element.append(t_element)
    p_element.append(r_element)
    return p_element


def insert_paragraphs_after_element(element, texts, style='Normal'):
    """Insert multiple paragraphs after a given XML element, preserving order."""
    for text in reversed(texts):
        p_element = create_new_paragraph_element(text, style)
        element.addnext(p_element)


def create_improved_document():
    # Cargar desde el documento YA MEJORADO por v2 (para no perder antecedentes, impactos, etc.)
    source_path = "D:/IX/detectorIA/documentacion/PROYECTO_INVESTIGACION_MEJORADO.docx"
    dest_path = "D:/IX/detectorIA/documentacion/PROYECTO_INVESTIGACION_MEJORADO.docx"

    doc = Document(source_path)
    cambios_count = 0

    print("=== MEJORANDO DOCUMENTO v4 ===")

    # ===================================================================
    # 1. MEJORAR METODOLOGIA - Agregar seccion de metodologia por objetivos
    # ===================================================================
    print("\n[1/4] Insertando Metodologia por Objetivos...")
    
    # Buscar "Tipo y enfoque de investigacion" como punto de referencia
    tipo_idx = None
    for i, p in enumerate(doc.paragraphs):
        if normalize("Tipo y enfoque de investigacion") in normalize(p.text or ''):
            tipo_idx = i
            break

    if tipo_idx:
        print(f"  'Tipo y enfoque' encontrado en indice {tipo_idx}")
        
        # Buscar "Diseno de investigacion" - insertaremos ANTES de esta seccion
        disenio_idx = None
        for j in range(tipo_idx + 1, len(doc.paragraphs)):
            if normalize("Diseno de investigacion") in normalize(doc.paragraphs[j].text or ''):
                disenio_idx = j
                break

        if disenio_idx:
            print(f"  'Diseno de investigacion' encontrado en indice {disenio_idx}")
            
            # Obtener lines de texto de METODOLOGIA_OBJETIVOS
            lines = [l.strip() for l in METODOLOGIA_OBJETIVOS.strip().split('\n') if l.strip()]
            
            # Insertar parrafos ANTES de "Diseno de investigacion"
            disenio_element = doc.paragraphs[disenio_idx]._element
            for line_text in reversed(lines):
                p_element = create_new_paragraph_element(line_text, 'Normal')
                disenio_element.addprevious(p_element)
            
            cambios_count += len(lines)
            print(f"  Insertados {len(lines)} parrafos de metodologia por objetivos")
        else:
            print("  ERROR: No se encontro 'Diseno de investigacion'")
    else:
        print("  ERROR: No se encontro 'Tipo y enfoque de investigacion'")

    # ===================================================================
    # 2. LIMPIAR SECCION RECURSOS - Reemplazar con contenido correcto
    # ===================================================================
    print("\n[2/4] Corrigiendo seccion Recursos...")
    
    recursos_start = None
    for i, p in enumerate(doc.paragraphs):
        if normalize("Recursos necesarios") in normalize(p.text or ''):
            recursos_start = i
            break

    if recursos_start:
        print(f"  'Recursos necesarios' encontrado en indice {recursos_start}")
        
        # Encontrar donde termina (proximo titulo de seccion)
        recursos_end = None
        for j in range(recursos_start + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text or ''
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ''
            if text.strip() and ('Cuerpo' in style or 'Heading' in style):
                recursos_end = j
                break
        if not recursos_end:
            recursos_end = len(doc.paragraphs)
        
        print(f"  Rango de recursos: {recursos_start + 1} a {recursos_end}")
        
        # Eliminar parrafos de recursos del XML (no solo limpiar texto)
        delete_paragraph_elements(doc, recursos_start + 1, recursos_end)
        
        # Agregar nuevo contenido como parrafos despues del titulo
        recursos_content = [
            "Para la ejecucion del presente proyecto de investigacion se requirieron los siguientes recursos, organizados en equipos de computo, software y datos.",
            "",
            "Equipos:",
            "- Una (1) laptop con procesador Intel Core i7 y 16 GB de RAM para el desarrollo local y pruebas iniciales.",
            "- Una (1) GPU Tesla T4 (14.6 GB VRAM) proporcionada por la plataforma Kaggle para el entrenamiento del modelo.",
            "",
            "Software y tecnologias:",
            "- Sistema operativo Windows 11 para el desarrollo local.",
            "- Python 3.12 como lenguaje de programacion principal.",
            "- PyTorch 2.x y Torchvision para la construccion y entrenamiento del modelo.",
            "- Bibliotecas auxiliares: scikit-learn, NumPy, Pandas, Matplotlib, Seaborn, OpenCV, Streamlit.",
            "- Dataset publico \"140k Real and Fake Faces\" de Kaggle para entrenamiento, validacion y prueba.",
            "- Plataforma Kaggle (con GPU) para el entrenamiento del modelo.",
            "",
            "Infraestructura:",
            "- Espacio de trabajo con computadora y conexion a internet de banda ancha.",
            "- Cuenta en Kaggle para acceso a datasets y GPU."
        ]
        
        insert_paragraphs_after_element(doc.paragraphs[recursos_start]._element, [t for t in recursos_content if t])
        cambios_count += len([t for t in recursos_content if t])
        print(f"  Insertados {len([t for t in recursos_content if t])} parrafos de recursos corregidos")
    else:
        print("  ERROR: No se encontro 'Recursos necesarios'")

    # ===================================================================
    # 3. ACTUALIZAR REFERENCIAS BIBLIOGRAFICAS
    # ===================================================================
    print("\n[3/4] Actualizando Referencias bibliograficas...")
    
    # Buscar la seccion de Referencias
    ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if normalize("Referencias") in normalize(p.text or '') and i > 100:
            ref_idx = i
            break

    if ref_idx is None:
        # Buscar al final del documento
        for i in range(len(doc.paragraphs) - 1, 0, -1):
            if normalize("Referencias") in normalize(doc.paragraphs[i].text or ''):
                ref_idx = i
                break

    if ref_idx:
        print(f"  'Referencias' encontrado en indice {ref_idx}")
        
        # Encontrar hasta donde llegan las referencias actuales
        ref_end = None
        for j in range(ref_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[j].text or ''
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ''
            if text.strip() and ('Cuerpo' in style or 'Heading' in style) and normalize('Referencias') not in normalize(text):
                ref_end = j
                break
        if not ref_end:
            ref_end = len(doc.paragraphs)
        
        # Eliminar referencias viejas del XML
        if ref_end > ref_idx + 1:
            delete_paragraph_elements(doc, ref_idx + 1, ref_end)
        
        # Nuevas referencias
        referencias = [
            "Arrieta, A. B., Diaz-Rodriguez, N., Del Ser, J., Bennetot, A., Tabik, S., Barbado, A., ... & Herrera, F. (2020). Explainable Artificial Intelligence (XAI): Concepts, taxonomies, opportunities and challenges toward responsible AI. Information Fusion, 58, 82-115.",
            "Brooke, J. (1996). SUS: A quick and dirty usability scale. Usability Evaluation in Industry, 189(194), 4-7.",
            "Flach, P. (2012). Machine Learning: The Art and Science of Algorithms that Make Sense of Data. Cambridge University Press.",
            "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
            "Hernandez-Sampieri, R. y Mendoza, C. (2018). Metodologia de la investigacion: Las rutas cuantitativa, cualitativa y mixta. McGraw-Hill Interamericana.",
            "Naupas, H., Valdivia, M., Palacios, J., y Romero, H. (2018). Metodologia de la investigacion cuantitativa-cualitativa y redaccion de la tesis (5a ed.). Ediciones de la U.",
            "Pan, S. J., & Yang, Q. (2010). A survey on transfer learning. IEEE Transactions on Knowledge and Data Engineering, 22(10), 1345-1359.",
            "Selvaraju, R. R., Cogswell, M., Das, A., Vedantam, R., Parikh, D., & Batra, D. (2017). Grad-CAM: Visual explanations from deep networks via gradient-based localization. Proceedings of the IEEE International Conference on Computer Vision (ICCV), 618-626.",
            "Sokolova, M., & Lapalme, G. (2009). A systematic analysis of performance measures for classification tasks. Information Processing & Management, 45(4), 427-437.",
            "Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley."
        ]
        
        # Insertar referencias despues del titulo
        insert_paragraphs_after_element(doc.paragraphs[ref_idx]._element, referencias)
        cambios_count += len(referencias)
        print(f"  Insertadas {len(referencias)} referencias bibliograficas")
    else:
        print("  ERROR: No se encontro la seccion 'Referencias'")

    # ===================================================================
    # 4. BARRIDO FINAL: Eliminar cualquier texto residual de IoT/Raspberry Pi
    # ===================================================================
    print("\n[4/4] Barrido final de contenido residual...")
    
    residual_keywords = [
        "Raspberry Pi", "raspberry", "IoT", "BoT-IoT", "TensorFlow Lite", 
        "TFLite", "TensorFlow 2.x", "Wireshark", "tcpdump",
        "microSD", "camara IP", "sensor de temperatura",
        "Raspberry Pi OS", "dispositivo IoT", "mesa de experimentacion"
    ]
    
    residual_encontrados = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text or ''
        for kw in residual_keywords:
            if kw.lower() in text.lower() and text.strip():
                # Limpiar este parrafo
                for run in p.runs:
                    run.text = ''
                if p.runs:
                    p.runs[0].text = ''
                else:
                    p.add_run('')
                residual_encontrados += 1
                print(f"  Limpiado [{i}]: contenido con '{kw}'")
                break

    if residual_encontrados == 0:
        print("  OK: Sin contenido residual de IoT/Raspberry Pi")
    else:
        print(f"  Limpiados {residual_encontrados} parrafos con contenido residual")

    # ===================================================================
    # 5. Guardar documento
    # ===================================================================
    doc.save(dest_path)
    print(f"\nDocumento guardado en: {dest_path}")
    print(f"Total de parrafos insertados/modificados: {cambios_count}")
    
    # ===================================================================
    # 6. VERIFICACION FINAL
    # ===================================================================
    print("\n=== VERIFICACION FINAL ===")
    
    # Verificar contenido residual
    verificar_kw = ["Raspberry", "IoT", "BoT-IoT", "TensorFlow Lite", "Wireshark"]
    errores = 0
    for i, p in enumerate(doc.paragraphs):
        for kw in verificar_kw:
            if kw.lower() in (p.text or '').lower():
                print(f"  ERROR: [{i}] Contiene '{kw}': {(p.text or '')[:100]}...")
                errores += 1
                break
    
    if errores == 0:
        print("  OK: Sin contenido residual")
    
    # Verificar metodologia por objetivos
    encontro_met = False
    for p in doc.paragraphs:
        if normalize("Objetivo Especifico 1") in normalize(p.text or ''):
            encontro_met = True
            break
    
    print(f"  Metodologia por Objetivos: {'OK' if encontro_met else 'NO ENCONTRADA'}")
    if not encontro_met:
        print("  Buscando texto similar...")
        for i, p in enumerate(doc.paragraphs):
            txt = normalize(p.text or '')
            if 'metodologia' in txt and 'objetivo' in txt:
                print(f"  Encontrado [{i}]: {(p.text or '')[:100]}...")
    
    # Verificar referencias
    ref_count = 0
    for p in doc.paragraphs:
        txt = normalize(p.text or '')
        if 'arxiv' in txt or 'ieee' in txt or 'mit press' in txt or 'mcgraw' in txt or 'addison' in txt:
            ref_count += 1
    
    print(f"  Referencias insertadas: {ref_count}")
    
    # Verificar recursos limpios
    recursos_ok = True
    for p in doc.paragraphs:
        txt = normalize(p.text or '')
        if 'raspberry' in txt or 'bot-iot' in txt or 'tensorflow lite' in txt:
            if p.text.strip():
                recursos_ok = False
                break
    
    print(f"  Recursos corregidos: {'OK' if recursos_ok else 'CONTENIDO RESIDUAL ENCONTRADO'}")

    return dest_path


if __name__ == "__main__":
    path = create_improved_document()
    print(f"\nArchivo creado: {path}")
