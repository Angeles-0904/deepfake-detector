"""
Genera el documento de ANEXOS en formato Word (.docx) para la tesis.
Incluye todas las imágenes generadas en la estructura académica solicitada.

Uso: python create_annexes_document.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ─── Rutas ────────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent
ANNEXES_DIR = PROJECT_ROOT / "annexes"

# Mapeo de imágenes por anexo
ANNEX_IMAGES = {
    "A": {
        "title": "Análisis de Características",
        "items": [
            ("A.1", "Comparación de histogramas de color RGB",
             "anexo_a/A1_histogramas_color.png",
             "Histogramas de color para los canales R, G, B en imágenes REAL y FAKE. "
             "Se observan diferencias en la distribución de intensidades, "
             "particularmente en el canal rojo."),
            ("A.2", "Detección de bordes con filtro Canny",
             "anexo_a/A2_deteccion_bordes_canny.png",
             "Aplicación del filtro Canny con diferentes umbrales para detectar "
             "bordes en imágenes REAL y FAKE. Las imágenes FAKE muestran "
             "patrones de borde menos definidos."),
            ("A.3", "Análisis de frecuencias con FFT",
             "anexo_a/A3_analisis_frecuencias_fft.png",
             "Transformada Rápida de Fourier (FFT) aplicada a imágenes REAL y FAKE. "
             "Se muestra el espectro de frecuencia y la reconstrucción con "
             "filtros pasa-bajos."),
        ]
    },
    "B": {
        "title": "Entrenamiento del Modelo",
        "items": [
            ("B.1", "Entorno de Kaggle con GPU Tesla T4",
             "anexo_b/B1_entorno_kaggle.png",
             "El entrenamiento se realizó en la plataforma Kaggle con "
             "aceleración GPU Tesla T4 (14.6 GB VRAM). Se utilizó "
             "PyTorch 2.x como framework de deep learning."),
            ("B.2", "Arquitectura del modelo DenseNet-121",
             "anexo_b/B2_arquitectura_densenet121.png",
             "Diagrama de la arquitectura DenseNet-121 adaptada para la "
             "clasificación binaria de deepfakes."),
            ("B.3", "Estrategia de congelamiento (Transfer Learning)",
             "anexo_b/B3_estrategia_congelamiento.png",
             "Estrategia de Transfer Learning: capas pre-entrenadas en "
             "ImageNet congeladas hasta Dense Block 3, fine-tuning en "
             "Dense Block 4 y clasificador."),
            ("B.4", "Curvas de entrenamiento (Loss y Accuracy)",
             "anexo_b/B4_curvas_entrenamiento.png",
             "Evolución de la pérdida y precisión durante las 14 épocas de "
             "entrenamiento. Se aplicó early stopping con paciencia de 7 épocas."),
            ("B.5", "Resumen del modelo y parámetros",
             "anexo_b/B5_resumen_modelo.png",
             "Resumen detallado del modelo mostrando todas las capas, "
             "parámetros totales (7,479,682) y entrenables (1,011,778)."),
        ]
    },
    "C": {
        "title": "Resultados de Evaluación",
        "items": [
            ("C.1", "Curvas de entrenamiento (Loss y Accuracy)",
             "anexo_c/C1_curvas_entrenamiento.png",
             "Las curvas de entrenamiento muestran convergencia estable del modelo, "
             "alcanzando accuracy de validación del 97.38%."),
            ("C.2", "Matriz de confusión",
             "anexo_c/C2_matriz_confusion.png",
             "Matriz de confusión en el conjunto de prueba (20,000 imágenes). "
             "VP ≈ 9,720, VN ≈ 9,740, FP ≈ 280, FN ≈ 260."),
            ("C.3", "Curva ROC",
             "anexo_c/C3_curva_roc.png",
             "Curva ROC con AUC = 0.9966, indicando una capacidad de "
             "discriminación casi perfecta entre clases REAL y FAKE."),
            ("C.4", "Tabla de métricas de clasificación",
             "anexo_c/C4_tabla_metricas.png",
             "Resumen de todas las métricas de evaluación: Accuracy 97.06%, "
             "Precision 0.97, Recall 0.97, F1-Score 0.97, AUC-ROC 0.9966."),
        ]
    },
    "D": {
        "title": "Mapas de Calor (Grad-CAM)",
        "items": [
            ("D.1", "Mapa de calor para imagen REAL",
             "anexo_d/D1_gradcam_real.png",
             "Grad-CAM aplicado a una imagen REAL. Las regiones de activación "
             "se distribuyen naturalmente sobre el rostro, con énfasis en "
             "ojos y centro facial."),
            ("D.2", "Mapa de calor para imagen FAKE",
             "anexo_d/D2_gradcam_fake.png",
             "Grad-CAM aplicado a una imagen FAKE. Se observan patrones de "
             "activación anómalos y dispersos, característicos de "
             "imágenes generadas por GANs."),
            ("D.3", "Explicaciones textuales generadas automáticamente",
             "anexo_d/D3_tabla_explicaciones.png",
             "Explicaciones textuales generadas por el sistema basadas en "
             "el análisis de las regiones de activación del Grad-CAM."),
        ]
    },
    "E": {
        "title": "Interfaz Web (Streamlit)",
        "items": [
            ("E.1", "Pantalla principal de la aplicación",
             "anexo_e/E1_pantalla_principal.png",
             "[CAPTURA REQUERIDA] Pantalla principal de la aplicación "
             "web desarrollada en Streamlit, mostrando el diseño y la "
             "interfaz de carga de imágenes."),
            ("E.2", "Resultado de clasificación para imagen REAL",
             "anexo_e/E2_resultado_real.png",
             "[CAPTURA REQUERIDA] Resultado de la clasificación para una "
             "imagen REAL, mostrando la predicción, nivel de confianza "
             "y métricas."),
            ("E.3", "Resultado de clasificación para imagen FAKE con heatmap",
             "anexo_e/E3_resultado_fake_heatmap.png",
             "[CAPTURA REQUERIDA] Resultado de clasificación para una "
             "imagen FAKE, incluyendo el mapa de calor Grad-CAM superpuesto."),
            ("E.4", "Explicación textual generada por el sistema",
             "anexo_e/E4_explicacion_textual.png",
             "[CAPTURA REQUERIDA] Sección de la interfaz que muestra "
             "la explicación textual generada automáticamente."),
            ("E.5", "Pruebas de robustez e historial de predicciones",
             "anexo_e/E5_robustez_historial.png",
             "[CAPTURA REQUERIDA] Sección de pruebas de robustez y "
             "historial de predicciones durante la sesión."),
        ]
    },
    "F": {
        "title": "Prueba de Compresión (Robustez)",
        "items": [
            ("F.1", "Tabla de precisión por nivel de compresión JPEG",
             "anexo_f/F1_tabla_compresion.png",
             "Tabla comparativa del rendimiento del modelo bajo diferentes "
             "niveles de compresión JPEG (QF=100, 75, 50)."),
            ("F.2", "Curva de robustez del modelo",
             "anexo_f/F2_curva_robustez.png",
             "Curva de robustez que muestra la degradación gradual del "
             "rendimiento al aumentar la compresión JPEG."),
        ]
    }
}


def create_document():
    """Crea el documento de Word con todos los anexos."""
    doc = Document()

    # ─── Configurar estilos ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ─── Portada de Anexos ────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ANEXOS")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Times New Roman"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Sistema de Detección de Deepfakes en Imágenes mediante\n"
        "Redes Neuronales Convolucionales con Transfer Learning\n"
        "e Inteligencia Artificial Explicable basada en Grad-CAM"
    )
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_paragraph()
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Bach. [Nombre del Autor]")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    doc.add_page_break()

    # ─── Generar cada anexo ───────────────────────────────────────────────
    for letter, annex in ANNEX_IMAGES.items():
        _add_annex_section(doc, letter, annex["title"], annex["items"])
        doc.add_page_break()

    # ─── Guardar documento ────────────────────────────────────────────────
    output_path = ANNEXES_DIR / "ANEXOS_DeepFake_Detector.docx"
    doc.save(str(output_path))
    print(f"\n[OK] Documento guardado en: {output_path}")
    return output_path


def _add_annex_section(doc, letter, title, items):
    """Agrega un anexo completo al documento."""
    # Título del anexo
    heading = doc.add_heading(f"Anexo {letter}: {title}", level=1)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)

    desc = doc.add_paragraph(
        f"En el presente anexo se presentan las evidencias correspondientes "
        f"al desarrollo del objetivo específico relacionado con {title.lower()}. "
        f"Las imágenes y tablas mostradas a continuación constituyen el "
        f"soporte visual de los resultados obtenidos durante la investigación."
    )
    desc.paragraph_format.first_line_indent = Cm(1.27)

    # Cada ítem del anexo
    for i, (code, name, img_path_rel, legend) in enumerate(items, 1):
        doc.add_paragraph()  # espacio

        # Subtítulo
        sub_heading = doc.add_heading(f"{code}: {name}", level=2)
        for run in sub_heading.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)

        # Imagen
        if img_path_rel:
            img_full_path = ANNEXES_DIR / img_path_rel
            if img_full_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_full_path), width=Inches(5.5))
            else:
                _add_placeholder(doc, code, name)
        else:
            _add_placeholder(doc, code, name)

        # Leyenda
        legend_p = doc.add_paragraph()
        legend_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = legend_p.add_run(
            f"{code}. {name}"
        )
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        # Descripción
        desc_p = doc.add_paragraph(legend)
        desc_p.paragraph_format.first_line_indent = Cm(1.27)
        desc_p.paragraph_format.space_after = Pt(6)
        for run in desc_p.runs:
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # Fuente
        fuente_p = doc.add_paragraph()
        fuente_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = fuente_p.add_run("Fuente: Elaboración propia.")
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"


def _add_placeholder(doc, code, name):
    """Agrega un placeholder para capturas pendientes o información textual."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"\n[ {code}: {name} ]\n"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(200, 0, 0)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "═══════════════════════════════════════════════\n"
        "     AQUÍ VA LA CAPTURA / EVIDENCIA VISUAL\n"
        "═══════════════════════════════════════════════\n"
    )
    run2.font.size = Pt(11)
    run2.font.name = "Courier New"


if __name__ == "__main__":
    print("=" * 60)
    print("  GENERANDO DOCUMENTO DE ANEXOS")
    print("=" * 60)

    path = create_document()

    print(f"\n[OK] Documento generado exitosamente.")
    print(f"[PATH] Ruta: {path}")
    print(f"\n[NOTA] Las capturas del Anexo E deben tomarse manualmente")
    print(f"   desde la app en Streamlit Cloud y reemplazar los placeholders.")
